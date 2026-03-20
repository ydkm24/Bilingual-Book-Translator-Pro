"""
export.py - PDF, DOCX, and ePub export methods.
Extracted from main.py during The Great Refactor as a mixin class.
"""
import os
import fitz
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import LETTER

from utils import get_app_path


class ExportMixin:
    """Mixin class providing all export methods (PDF, DOCX, ePub)."""

    def _save_visual_pdf(self, file_path):
        """Advanced export: Overlays translated text onto the original page images."""
        try:
            out_doc = fitz.open()
            # SPRINT 100: Aggressive page count recovery for exports
            pdf_path = self.current_pdf_path if self.current_pdf_path else self.get_valid_pdf_path()
            if not pdf_path: 
                messagebox.showerror("Export Failed", "Cannot locate original PDF to determine page count.")
                return
            
            ui_doc = self.ui_pdf_doc if self.ui_pdf_doc else fitz.open(pdf_path)
            num_pages = ui_doc.page_count
            
            # Synchronize truth back to memory
            if len(self.all_page_data) < num_pages:
                self.all_page_data.extend([{"page": i} for i in range(len(self.all_page_data), num_pages)])
            self.total_pages = num_pages
            
            for i in range(num_pages):
                # Search for data for this page (Robust lookup)
                data = self.all_page_data[i] if (i < len(self.all_page_data) and self.all_page_data[i]) else {"page": i}
                
                # Get translation
                trans_text = str(data.get('english', '') or '')
                
                # 1. Background (Original Image)
                img_path = data.get('cover_image_path')
                if not img_path and getattr(self, "cache_dir", None):
                    test_path = os.path.join(self.cache_dir, f"img_{i}.jpg")
                    if os.path.exists(test_path): img_path = test_path
                
                if img_path and os.path.exists(img_path):
                    img = Image.open(img_path)
                    w, h = img.size
                    page = out_doc.new_page(width=w, height=h)
                    page.insert_image(page.rect, filename=img_path)
                    
                    # 2. Text Overlay
                    blocks = data.get('blocks', [])
                    if blocks and trans_text and "[Translation Pending]" not in trans_text:
                        # Scale font based on width
                        scaled_size = max(10, int(page.rect.width / 45))
                        # Readability Box: Semi-transparent white background
                        page.insert_textbox(page.rect, trans_text, color=(0, 0, 0), fontsize=scaled_size, 
                                         fontname="helv", align=fitz.TEXT_ALIGN_LEFT, 
                                         overlay=True, fill_opacity=0.85)
                else:
                    # Fallback for pages without images/data
                    page = out_doc.new_page()
                    display_text = trans_text if trans_text and "[Translation Pending]" not in trans_text else "[Processing...]"
                    page.insert_text((50, 50), f"Page {i+1}\n\n" + display_text)

            # SPRINT 98: High-efficiency PDF save (Garbage collection + Deflate)
            out_doc.save(file_path, garbage=3, deflate=True)
            out_doc.close()
            messagebox.showinfo("Success", "Visual Fidelity PDF exported!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Visual PDF failed: {e}")

    def _save_standard_pdf(self, file_path):
        """Standard ReportLab bilingual table export."""
        try:
            doc = SimpleDocTemplate(file_path, pagesize=LETTER, 
                                    rightMargin=50, leftMargin=50, 
                                    topMargin=50, bottomMargin=50)
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle('TitleStyle', parent=styles['Heading1'], fontSize=24, spaceAfter=20, alignment=1)
            column_style = ParagraphStyle('ColumnStyle', parent=styles['Normal'], fontSize=10, leading=12)
            
            story = []
            title_text = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Translation"
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 0.5 * inch))
            
            pdf_path = self.current_pdf_path if self.current_pdf_path else self.get_valid_pdf_path()
            if not pdf_path: 
                messagebox.showerror("Export Failed", "Cannot locate original PDF.")
                return
            ui_doc = self.ui_pdf_doc if self.ui_pdf_doc else fitz.open(pdf_path)
            num_pages = ui_doc.page_count
            
            # Synchronize truth back to memory
            if len(self.all_page_data) < num_pages:
                self.all_page_data.extend([{"page": i} for i in range(len(self.all_page_data), num_pages)])
            self.total_pages = num_pages # Update truth
            for i in range(num_pages):
                # Robust data lookup
                data = self.all_page_data[i] if (i < len(self.all_page_data) and self.all_page_data[i]) else {"page": i}
                story.append(Paragraph(f"Page {i + 1}", styles['Heading2']))
                story.append(Spacer(1, 0.1 * inch))
                
                orig_text = str(data.get('original', '')).replace('\n', '<br/>')
                
                # Get translation
                trans_text = str(data.get('english', '') or '')
                
                
                trans_text = trans_text.replace('\n', '<br/>')
                
                # Professional alignment for side-by-side reading
                orig_lines = orig_text.split('<br/>')
                trans_lines = trans_text.split('<br/>')
                max_lines = max(len(orig_lines), len(trans_lines))
                
                # Process in chunks of 15 lines per table row to avoid ReportLab overflow
                chunk_size = 15
                for idx_chunk in range(0, max_lines, chunk_size):
                    o_chunk = "<br/>".join(orig_lines[idx_chunk:idx_chunk+chunk_size])
                    t_chunk = "<br/>".join(trans_lines[idx_chunk:idx_chunk+chunk_size])
                    
                    row_content = []
                    # Find image path
                    curr_img_path = data.get('cover_image_path')
                    if not curr_img_path and getattr(self, "cache_dir", None):
                        test_path = os.path.join(self.cache_dir, f"img_{i}.jpg")
                        if os.path.exists(test_path): curr_img_path = test_path

                    if idx_chunk == 0 and curr_img_path and os.path.exists(curr_img_path):
                        # SPRINT 98: Strict Image Scaling to avoid ReportLab overflow
                        with Image.open(curr_img_path) as pil_img:
                            orig_w, orig_h = pil_img.size
                        
                        max_w = 2.8 * inch
                        max_h = 4.0 * inch
                        # Compute scale to fit within both max_w and max_h
                        scale = min(max_w / orig_w, max_h / orig_h)
                        draw_w = orig_w * scale
                        draw_h = orig_h * scale
                        
                        img = RLImage(curr_img_path, width=draw_w, height=draw_h)
                        # CRITICAL: Force draw dimensions to prevent ReportLab ignoring logic
                        img.drawWidth = draw_w
                        img.drawHeight = draw_h
                        row_content.append([img, Paragraph(o_chunk, column_style)])
                    else:
                        row_content.append(Paragraph(o_chunk, column_style))
                    
                    row_content.append(Paragraph(t_chunk, column_style))
                    
                    table_data = [[row_content[0], row_content[1]]]
                    t = Table(table_data, colWidths=[3 * inch, 3 * inch])
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'), 
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey), 
                        ('TOPPADDING', (0,0), (-1,-1), 12)
                    ]))
                    story.append(t)
                story.append(Spacer(1, 0.3 * inch))

            doc.build(story)
            messagebox.showinfo("Success", "Professional Bilingual PDF exported!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to build PDF: {str(e)}")

    def _save_facing_pdf(self, file_path):
        """Standard 'Bilingual Book' format: Original page on left, Translated on right."""
        try:
            out_doc = fitz.open()
            pdf_path = self.current_pdf_path if self.current_pdf_path else self.get_valid_pdf_path()
            if not pdf_path:
                messagebox.showerror("Export Failed", "Cannot locate original PDF.")
                return
            ui_doc = self.ui_pdf_doc if self.ui_pdf_doc else fitz.open(pdf_path)
            num_pages = ui_doc.page_count
            
            # Synchronize truth back to memory
            if len(self.all_page_data) < num_pages:
                self.all_page_data.extend([{"page": i} for i in range(len(self.all_page_data), num_pages)])
            self.total_pages = num_pages
            
            for i in range(num_pages):
                # Robust data lookup
                data = self.all_page_data[i] if (i < len(self.all_page_data) and self.all_page_data[i]) else {"page": i}
                
                # Get translation
                trans_text = str(data.get('english', '') or '')

                # 1. Left Page (Original)
                img_path = data.get('cover_image_path')
                if not img_path and getattr(self, "cache_dir", None):
                    test_path = os.path.join(self.cache_dir, f"img_{i}.jpg")
                    if os.path.exists(test_path): img_path = test_path

                if img_path and os.path.exists(img_path):
                    img = Image.open(img_path)
                    w, h = img.size
                    page_orig = out_doc.new_page(width=w, height=h)
                    page_orig.insert_image(page_orig.rect, filename=img_path)
                else:
                    page_orig = out_doc.new_page()
                    page_orig.insert_text((50, 50), f"Page {i + 1} (Original)\n\n" + str(data.get('original', '')))
                
                # 2. Right Page (Translation)
                # Matches the same dimensions for consistent reading flow
                w_t = page_orig.rect.width
                h_t = page_orig.rect.height
                page_trans = out_doc.new_page(width=w_t, height=h_t)
                
                
                if not trans_text: trans_text = "[No translation available for this page yet]"

                scaled_size = max(11, int(w_t / 40)) # Slightly larger for dedicated page
                
                # Use a standard margin for the translation page
                margin = 50
                rect_trans = fitz.Rect(margin, margin, w_t - margin, h_t - margin)
                page_trans.insert_textbox(rect_trans, trans_text, color=(0,0,0), 
                                         fontsize=scaled_size, fontname="helv", align=fitz.TEXT_ALIGN_LEFT)

            out_doc.save(file_path, garbage=3, deflate=True)
            out_doc.close()
            messagebox.showinfo("Success", "Facing Pages PDF exported!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Facing PDF failed: {e}")

    def handle_export(self, choice):
        if choice == "Export PDF":
            self.save_as_pdf()
        elif choice == "Export Word":
            self.save_as_docx()
        elif choice == "Export ePub":
            self.save_as_epub()

    def save_as_epub(self):
        """Exports the book to a professional paragraph-toggle ePub format."""
        try:
            from ebooklib import epub
        except ImportError:
            messagebox.showwarning("Missing Library", "Please install 'ebooklib' and 'lxml' to use this feature.")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".epub", filetypes=[("ePub Files", "*.epub")])
        if not file_path: return

        try:
            book = epub.EpubBook()
            title = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Book"
            book.set_title(title)
            book.set_language('en')
            
            # Add basic CSS for bilingual layout
            style = 'body { font-family: sans-serif; } .pair { margin-bottom: 2em; border-left: 4px solid #4A90E2; padding-left: 10px; } .orig { color: #666; font-style: italic; margin-bottom: 5px; } .trans { font-weight: bold; color: #333; }'
            nav_css = epub.EpubItem(uid="style_nav", file_name="style/nav.css", media_type="text/css", content=style)
            book.add_item(nav_css)

            chapters = []
            for data in self.all_page_data:
                p_num = data['page'] + 1
                c = epub.EpubHtml(title=f"Page {p_num}", file_name=f"page_{p_num}.xhtml", lang='en')
                c.add_item(nav_css)
                
                content = f"<h1>Page {p_num}</h1>"
                
                # Split original and translation into paragraphs for pairing
                orig_paras = [p.strip() for p in data.get('original', '').split('\n') if p.strip()]
                trans_paras = [p.strip() for p in data.get('english', '').split('\n') if p.strip()]
                
                # Zip them together (best effort alignment)
                import itertools
                for o, t in itertools.zip_longest(orig_paras, trans_paras, fillvalue=""):
                    content += f'<div class="pair"><div class="orig">{o}</div><div class="trans">{t}</div></div>'
                
                c.content = content
                book.add_item(c)
                chapters.append(c)

            book.toc = tuple(chapters)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav'] + chapters
            epub.write_epub(file_path, book, {})
            messagebox.showinfo("Success", "Paragraph-Toggle ePub exported successfully!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to save ePub: {str(e)}")


    def save_as_pdf(self):
        """Exports the book to PDF with multiple layout options."""
        dialog = ctk.CTkToplevel(self)
        dialog.title("Export PDF")
        dialog.geometry("400x300")
        dialog.transient(self)
        dialog.grab_set()

        ctk.CTkLabel(dialog, text="Choose Export Layout:", font=("Inter", 16, "bold")).pack(pady=20)
        
        def _choice(mode):
            dialog.destroy()
            file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
            if not file_path: return
            
            if mode == "visual":
                self._save_visual_pdf(file_path)
            elif mode == "facing":
                self._save_facing_pdf(file_path)
            else:
                self._save_standard_pdf(file_path)

        ctk.CTkButton(dialog, text="Visual Overlay (Fidelity)", command=lambda: _choice("visual")).pack(pady=5, padx=20, fill="x")
        ctk.CTkButton(dialog, text="Facing Pages (Left: Orig, Right: Trans)", command=lambda: _choice("facing")).pack(pady=5, padx=20, fill="x")
        ctk.CTkButton(dialog, text="Standard Bilingual (Side-by-Side Table)", command=lambda: _choice("standard")).pack(pady=5, padx=20, fill="x")
        ctk.CTkButton(dialog, text="Cancel", fg_color="gray", command=dialog.destroy).pack(pady=20)

    def save_as_docx(self):
        """Exports the Bilingual translation to a Word document."""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            messagebox.showwarning("Missing Library", "Please install 'python-docx' to use this feature.\nRun: pip install python-docx")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
        if not file_path: return

        try:
            doc = Document()
            title = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Translation"
            doc.add_heading(title, 0)

            for data in self.all_page_data:
                doc.add_heading(f"Page {data['page'] + 1}", level=1)
                
                # Table for side-by-side
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                cells = table.rows[0].cells
                
                cells[0].text = data.get('original', '')
                cells[1].text = data.get('english', '')
                
                # Image if it exists
                if data.get('is_image') and data.get('cover_image_path') and os.path.exists(data['cover_image_path']):
                    doc.add_picture(data['cover_image_path'], width=Inches(4))

            doc.save(file_path)
            messagebox.showinfo("Success", "Word document exported successfully!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to save Word doc: {str(e)}")

