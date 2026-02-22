import os
import sys
import threading
import time
import random
import multiprocessing
import concurrent.futures
from functools import partial
from typing import Any, Dict, List, Optional, cast
import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import fitz  # PyMuPDF
from deep_translator import GoogleTranslator
import pytesseract
from PIL import Image, ImageOps, ImageFilter, ImageTk
import io
import re
import cv2
import json
import numpy as np
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.pagesizes import LETTER
from supabase import create_client, Client
from dotenv import load_dotenv

# --- SPRINT 37: GLOBAL CONFIG & PORTABILITY ---
APP_VERSION = "1.0.4"

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    try:
        # PyInstaller creates a temp folder and stores path in _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

def get_app_path(sub_path=""):
    """Get path to the user's AppData/Local folder for the app."""
    path = os.path.join(os.environ.get('LOCALAPPDATA', os.path.expanduser('~')), "PDF_Translator")
    if sub_path:
        path = os.path.join(path, sub_path)
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)
    return path

# Load Supabase credentials
if getattr(sys, 'frozen', False):
    # If built with PyInstaller, use the bundled .env
    env_path = os.path.join(sys._MEIPASS, ".env")
    load_dotenv(env_path)
else:
    load_dotenv()

# Set appearance mode and color theme
ctk.set_appearance_mode("Dark") 
ctk.set_default_color_theme("blue")

# True Light / Midnight Pro Palette Tuples (Light, Dark)
PAPER_BG = ("#FDF5E6", "#121212") # Beige / Deep matte black
TEXT_COLOR = ("#111111", "#E0E0E0") # Primary text
TEXT_DIM = ("#555555", "#A0A0A0") # Secondary text
HEADER_BG = ("#F5E6CC", "#1E1E1E") # Elevated frame background
ACCENT_COLOR = ("#0056b3", "#00D2FF") # Electric Blue
ACCENT_HOVER = ("#004494", "#00B4DB")
PAPER_SHEET_BG = ("#FFFFFF", "#242424") # Surface color for pages
BORDER_COLOR = ("#D4C4A8", "#333333") 
CARD_BG = ("#E8D5B5", "#2D2D2D")
FRAME_BG = ("#F5E6CC", "#1E1E1E")
BTN_PRIMARY = ("#6200EE", "#BB86FC")
BTN_PRIMARY_HOVER = ("#3700B3", "#9965f4")
BTN_SECONDARY = ("#E0E0E0", "#333333")
BTN_SECONDARY_HOVER = ("#d5c09e", "#444444")
INPUT_BG = ("#FFFFFF", "#2D2D2D")
BTN_DANGER = ("#E53935", "#C62828")
BTN_DANGER_HOVER = ("#D32F2F", "#B71C1C")
BTN_SUCCESS = ("#4CAF50", "#27AE60")
BTN_SUCCESS_HOVER = ("#45A049", "#2ECC71")
BTN_WARNING = ("#FFB74D", "#FF9800")
BTN_WARNING_HOVER = ("#FFA726", "#F57C00")

def ocr_worker(page_num, file_path, ocr_lang_code, is_auto, languages, translator_src, page_width, current_rtl, tesseract_path):
    """Standalone worker for processing a single page in a separate process.
    Reads image directly from file and extracts layout blocks independently."""
    # We re-import inside the worker to ensure process-isolation safety
    import pytesseract
    from PIL import Image
    import io
    from deep_translator import GoogleTranslator
    import re
    import cv2
    import numpy as np
    import fitz
    
    # CRITICAL: Path must be set PER PROCESS on Windows. Resolved via Phase A abstraction.
    pytesseract.pytesseract.tesseract_cmd = tesseract_path
    
    # Helper functions copied/needed inside worker
    def clean_ocr_text(text):
        if not text: return ""
        # Preserve double newlines for paragraph structure
        text = re.sub(r'[|\\/_]', '', text)
        segments = text.split('\n\n')
        cleaned_segments = [re.sub(r'\s+', ' ', s).strip() for s in segments]
        return '\n\n'.join(cleaned_segments).strip()

    def detect_script(text, languages):
        text = str(text)
        is_rtl = False
        detected_ocr = []
        detected_trans = "auto"
        
        # Check for scripts
        if any("\u0590" <= c <= "\u05FF" for c in text): # Hebrew
            detected_ocr.append("heb")
            is_rtl = True
            detected_trans = "iw"
        if any("\u0600" <= c <= "\u06FF" for c in text): # Arabic
            detected_ocr.append("ara")
            is_rtl = True
            detected_trans = "ar"
        if any("\u0370" <= c <= "\u03FF" for c in text): # Greek
            detected_ocr.append("ell")
            detected_trans = "el"
        if any("\u0400" <= c <= "\u04FF" for c in text): # Cyrillic
            detected_ocr.append("rus")
            detected_trans = "ru"
        if any("\u4e00" <= c <= "\u9fff" for c in text): # Chinese
            detected_ocr.append("chi_sim")
            detected_trans = "zh-CN"
            
        if not detected_ocr:
            return {"ocr": "fra+eng", "trans": "fr", "rtl": False}
            
        return {"ocr": "+".join(detected_ocr), "trans": detected_trans, "rtl": is_rtl}

    doc = fitz.open(file_path)
    page = doc.load_page(page_num)
    
    # Extract blocks inside the worker to prevent main-thread hangs
    blocks = list(page.get_text("blocks"))
    
    # Using 200 DPI (Matrix 2,2) for standard workers to save RAM
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img_bytes = pix.tobytes("png")
    img = Image.open(io.BytesIO(img_bytes))
    doc.close() # Release handle immediately
    
    # ... (Logic from pass 1 and pass 2, including preprocessing and translation)
    # This is a full extraction of the page processing logic to be independent of self
    try:
        # Preprocessing
        cv_img = cv2.imdecode(np.frombuffer(img_bytes, np.uint8), cv2.IMREAD_COLOR)
        if cv_img is None: return {"page": page_num, "error": "Image decode failed"}
        
        # SPRINT 31: Advanced Preprocessing (Speckle Filter & Contrast for OCR Worker)
        # We use a fast version for the worker
        gray = cv2.cvtColor(cv_img, cv2.COLOR_BGR2GRAY)
        denoised = cv2.medianBlur(gray, 3)
        thresh = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # SPRINT 30/31: Speckle Filter (Dust Removal) inside Worker
        labels_info = cv2.connectedComponentsWithStats(cv2.bitwise_not(thresh), connectivity=8)
        num_labels, labels, stats, _ = labels_info
        for j in range(1, num_labels):
            if stats[j, cv2.CC_STAT_AREA] < 10: thresh[labels == j] = 255
            
        ink_density = np.sum(thresh == 0) / thresh.size
        processed_pil = Image.fromarray(thresh)
        
        # Initial state
        translated_text = ""
        literal_text = ""
        ocr_text = "" 
        page_is_rtl = current_rtl if not is_auto else False
        ocr_lang = ocr_lang_code if not is_auto else "fra+eng"
        is_real_image = False

        # SPRINT 30: Blank Page Detection
        if ink_density < 0.0005:
            is_real_image = True
            ocr_text = "[No Text]"
        else:
            if is_auto:
                raw_detect = pytesseract.image_to_string(processed_pil, lang="eng+ara+heb+ell+rus+chi_sim", config='--oem 1 --psm 3').strip()
                detected = detect_script(raw_detect, languages)
                ocr_lang = detected["ocr"]
                page_is_rtl = detected["rtl"]
                translator_src = detected["trans"]
            
            # SPRINT 29/30: PSM 1 for Arabic OCR
            psm_cfg = '--oem 1 --psm 1' if "ara" in ocr_lang else '--oem 1 --psm 3'
            raw_ocr = pytesseract.image_to_string(processed_pil, lang=ocr_lang, config=psm_cfg).strip()
            ocr_text = clean_ocr_text(raw_ocr)
            
            # Phase 1 extraction cleanup
            if len(ocr_text) < 15: is_real_image = True
            
        # SPRINT 31: Numeral Shield for original text
        ocr_text = re.sub(r'(\d+[\-/]\d+[\-/]?[0-9]*[\-/]?[0-9]*)', u' \u2066\\1\u2069 ', ocr_text)
        
        # Phase 1: We only do Extraction now. Translation is removed from the process-level worker.

        # Detect is_image/nonsense
        is_nonsense = False
        alnum_count = 0
        if ocr_text:
            alnum_count = len(re.findall(r'[\w]', ocr_text))
            if alnum_count < len(ocr_text) * 0.3: is_nonsense = True
            
        # Fallback Logic: In Phase 1 we just ensure we have clean ocr_text.
        # Translation will be handled by Phase 2 in the main app.
        
        # Sensitivity Boost: Titles can be short (Pic 3), so we relax thresholds
        is_real_image = (len(ocr_text) < 10 and alnum_count < 5) or is_nonsense

        # Save the original image to cache for UI and Cloud Uploads
        img_path = None
        try:
            cache_dir = os.path.join(os.getcwd(), ".translator_cache", os.path.basename(file_path).replace(".pdf", ""))
            os.makedirs(cache_dir, exist_ok=True)
            img_path = os.path.join(cache_dir, f"img_{page_num}.png")
            img.save(img_path)
        except Exception as e:
            print(f"Failed to save image locally: {e}")

        print(f"[Worker] Finished Page {page_num} extraction.")
        return {
            "page": page_num,
            "original": ocr_text if ocr_text else "[No Text Detected]",
            "english": "[Translation Pending]",
            "literal": "[Loading...]",
            "is_image": is_real_image,
            "is_rtl_page": page_is_rtl,
            "translator_src": translator_src,
            "is_centered": (page_num == 0),
            "is_cover": (page_num == 0),
            "cover_image_path": img_path
        }
    except Exception as e:
        return {"page": page_num, "error": str(e)}

def translate_worker(text, translator_src, engine="Google", deepl_key="", openai_key="", glossary=None):
    """Standalone worker for translating text with robust retries, throttling, and glossary enforcement."""
    from deep_translator import GoogleTranslator
    import time
    import random
    import re
    
    # STABILITY FIX: Mandatory staggered launch to avoid "Burst" blocking
    time.sleep(random.uniform(0.1, 0.5))
    
    # Pre-processing Glossary for standard engines (Google/DeepL)
    if glossary and engine != "GPT-4o":
        for orig, target in glossary.items():
            # Use regex to replace whole words only to avoid partial matches
            pattern = re.compile(re.escape(orig), re.IGNORECASE)
            text = pattern.sub(target, text)
            
    last_err = ""
    # Try 3 times with exponential backoff
    for attempt in range(3):
        try:
            # SPRINT 31: Unicode Isolation Shield for Numbers
            text = re.sub(r'(\d[\-/]\d[\-/]?[0-9]*[\-/]?[0-9]*)', u' \u2066\\1\u2069 ', text)
            
            translated = "[Translation Error]"
            if engine == "DeepL" and deepl_key:
                import deepl
                translator = deepl.Translator(deepl_key)
                src = None if translator_src == 'auto' else translator_src.upper()
                
                # Contextual Flow: Split into paragraphs and translate as a batch
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                if paragraphs:
                    res = translator.translate_text(paragraphs, target_lang="EN-US", source_lang=src)
                    # res is a list of results if a list was passed
                    translated = "\n\n".join([r.text for r in res])
                else:
                    translated = ""
            elif engine == "GPT-4o" and openai_key:
                from openai import OpenAI
                import json as pyjson
                client = OpenAI(api_key=openai_key)
                
                # Split text into paragraphs for structured translation
                paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
                
                gloss_instr = ""
                if glossary:
                    gloss_list = [f"'{k}' -> '{v}'" for k, v in glossary.items()]
                    gloss_instr = f"IMPORTANT: Follow this glossary: {', '.join(gloss_list)}."
                
                prompt = (
                    f"Translate these {len(paragraphs)} paragraphs from {translator_src} to English. "
                    "Maintain the exact order and nuances. Use the provided glossary rules. "
                    "Return ONLY a JSON object with a key 'translations' containing an array of strings."
                )
                
                response = client.chat.completions.create(
                  model="gpt-4o",
                  response_format={ "type": "json_object" },
                  messages=[
                    {"role": "system", "content": f"You are a professional literary translator. {gloss_instr}"},
                    {"role": "user", "content": f"{prompt}\n\nTEXT TO TRANSLATE:\n" + pyjson.dumps(paragraphs)}
                  ]
                )
                
                try:
                    res_raw = response.choices[0].message.content
                    res_data = pyjson.loads(res_raw)
                    translated_list = res_data.get("translations", [])
                    if len(translated_list) == len(paragraphs):
                        translated = "\n\n".join(translated_list)
                    else:
                        # Fallback if AI gets the count wrong
                        translated = "\n\n".join(translated_list) if translated_list else res_raw
                except:
                    translated = response.choices[0].message.content.strip()
                    
            else: # Default Google
                translator = GoogleTranslator(source=translator_src, target='en')
                translated = translator.translate(text)
            
            words = text.split()
            literal = " ".join(words[:50]) + ("..." if len(words) > 50 else "")
            
            return {"translated": translated, "literal": literal}
        except Exception as e:
            last_err = str(e)
            time.sleep(2 ** attempt + random.uniform(0.1, 0.5))
            
    return {"error": f"Failed after 3 retries: {last_err}"}

class PDFTranslatorApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.title("Bilingual Book Translator Pro")
        self.geometry("1450x900")
        self.configure(fg_color=PAPER_BG)

        # --- Base State & Configuration ---
        self.config_file = get_app_path("config.json")
        self.selected_engine = "Google" # Default
        self.deepl_key = ""
        self.openai_key = ""
        self.glossary = {} # Original -> Target map
        self.load_settings()
        
        # --- Language Mapping ---
        self.languages = {
            "Spanish": {"ocr": "spa", "trans": "es"},
            "French": {"ocr": "fra", "trans": "fr"},
            "German": {"ocr": "deu", "trans": "de"},
            "Italian": {"ocr": "ita", "trans": "it"},
            "Portuguese": {"ocr": "por", "trans": "pt"},
            "Russian": {"ocr": "rus", "trans": "ru"},
            "Chinese": {"ocr": "chi_sim", "trans": "zh-CN"},
            "Arabic": {"ocr": "ara", "trans": "ar"},
            "Greek": {"ocr": "ell", "trans": "el"},
            "Latin": {"ocr": "lat", "trans": "la"},
            "Hebrew": {"ocr": "heb", "trans": "he"},
            "Yiddish": {"ocr": "yid", "trans": "yi"},
            "Auto-Detect": {"ocr": "eng", "trans": "auto"}
        }

        # Boolean Vars for Menu (Toggles)
        self.one_one_var = ctk.BooleanVar(value=False)
        self.deep_cleanup_var = ctk.BooleanVar(value=False)
        self.speed_var = ctk.BooleanVar(value=False)
        self.eco_var = ctk.BooleanVar(value=False)

        # --- Supabase Initialization & Auth ---
        self.supabase_url = os.getenv("SUPABASE_URL")
        self.supabase_key = os.getenv("SUPABASE_KEY")
        self.supabase: Client = None
        self.current_user = None # type: Optional[Dict[str, Any]]
        self.current_username = None # type: Optional[str]
        
        if self.supabase_url and self.supabase_key and isinstance(self.supabase_url, str) and "YOUR_SUPABASE" not in self.supabase_url:
            try:
                self.supabase = create_client(self.supabase_url, self.supabase_key)
                
                # Try to restore session
                session_data = self.load_auth_session()
                if session_data and session_data.get('access_token'):
                    try:
                        res = self.supabase.auth.set_session(session_data['access_token'], session_data['refresh_token'])
                        if res and res.user:
                            self.current_user = {"id": res.user.id, "email": res.user.email}
                            self.fetch_username()
                    except Exception as e:
                        print(f"Session restore failed: {e}")
                        
            except Exception as e:
                print(f"Supabase Init Error: {e}")
        
        # --- OCR Configuration ---
        self.tesseract_path = resource_path(r"bin\tesseract\tesseract.exe")
        # For development, check if it exists in local bin first, else fallback to hardcoded (safety)
        if not os.path.exists(self.tesseract_path):
             self.tesseract_path = r'E:\Tesseract-OCR\tesseract.exe'
             
        pytesseract.pytesseract.tesseract_cmd = self.tesseract_path

        # Grid layout configuration
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(2, weight=1) 
        
        self.active_executor = None # type: Any
        self.translation_executor = concurrent.futures.ThreadPoolExecutor(max_workers=5)
        self.ui_pdf_doc = None # Persistent handle for the UI thread
        self.current_session = 0 # Track session ID to isolate background tasks
        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.grid_rowconfigure(0, weight=0) # Controls
        self.grid_rowconfigure(1, weight=0) # Headers
        self.grid_rowconfigure(2, weight=1) # Content
        
        # Top Frame for controls (Decluttered)
        self.control_frame = ctk.CTkFrame(self, fg_color=HEADER_BG, corner_radius=0)
        self.control_frame.grid(row=0, column=0, sticky="ew")
        
        self.lang_label = ctk.CTkLabel(self.control_frame, text="Source Language:", text_color=TEXT_COLOR, font=("Arial", 12, "bold"))
        self.lang_label.pack(side="left", padx=(10, 5), pady=10)
        
        self.lang_menu = ctk.CTkOptionMenu(self.control_frame, values=list(self.languages.keys()), 
                                           fg_color=CARD_BG, button_color=CARD_BG, button_hover_color=BTN_SECONDARY_HOVER, width=120)
        self.lang_menu.pack(side="left", padx=5, pady=10)
        self.lang_menu.set("Auto-Detect")
        
        self.progress_bar = ctk.CTkProgressBar(self.control_frame, width=300, progress_color=ACCENT_COLOR)
        self.progress_bar.pack(side="left", padx=20, pady=10)
        self.progress_bar.set(0)
        
        self.status_label = ctk.CTkLabel(self.control_frame, text="Waiting for PDF...", text_color=TEXT_COLOR)
        self.status_label.pack(side="left", padx=10, pady=10)

        # Avatar / Login Button (Far Right)
        self.avatar_btn = ctk.CTkButton(self.control_frame, text="👤 Login", width=80,
                                        fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER, text_color=TEXT_COLOR,
                                        command=self.handle_avatar_click)
        self.avatar_btn.pack(side="right", padx=10, pady=10)

        # Search bar (Right)
        self.search_var = ctk.StringVar()
        self.search_entry = ctk.CTkEntry(self.control_frame, placeholder_text="Search book...", 
                                         textvariable=self.search_var, width=150)
        self.search_entry.pack(side="right", padx=10, pady=10)
        self.search_entry.bind("<Return>", self.search_text)        

        # Action Buttons (Keep primary ones visible)
        self.sync_btn = ctk.CTkButton(self.control_frame, text="⇪ Sync", command=self.force_cloud_sync,
                                          width=60, fg_color=("#9B59B6", "#8E44AD"), hover_color=("#8E44AD", "#9B59B6"), state="disabled", text_color="#FFFFFF")
        # Note: sync_btn is NOT packed by default. It appears only for cloud books.

        self.publish_btn = ctk.CTkButton(self.control_frame, text="✨ Publish", command=self.publish_to_gallery,
                                          width=80, fg_color=("#FFB74D", "#FF9800"), hover_color=("#FFA726", "#F57C00"), state="disabled", text_color="#121212")
        self.publish_btn.pack(side="right", padx=5, pady=10)

        self.pause_btn = ctk.CTkButton(self.control_frame, text="Pause", command=self.toggle_pause,
                                      fg_color=("#E53935", "#C62828"), hover_color=("#D32F2F", "#B71C1C"), width=60, state="disabled", text_color="#FFFFFF")
        self.pause_btn.pack(side="right", padx=(5, 10), pady=10)
        
        # Settings Menu Dropdown Button
        self.settings_btn = ctk.CTkButton(self.control_frame, text="⚙", width=30, command=self.toggle_settings_popdown, fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER, text_color=TEXT_COLOR)
        self.settings_btn.pack(side="right", padx=5, pady=10)

        self.settings_popdown = None

        # --- UI: Top Menu Bar ---
        self.create_menu_bar()

        if self.supabase:
            self.publish_btn.configure(state="normal")
            self.sync_btn.pack_forget() # Hide Sync for new local upload
        else:
            self.publish_btn.configure(state="disabled")
            self.sync_btn.configure(state="disabled")
            
        # SPRINT 36.5: Sync the Top-Right UI with the pre-loaded authed session if one exists
        if hasattr(self, 'avatar_btn'):
            self.update_avatar_ui()
            
        # SPRINT 37: Check for updates on startup
        threading.Thread(target=self.check_for_updates, daemon=True).start()

    def check_for_updates(self):
        """Checks Supabase Storage for a newer manifest.json."""
        if not self.supabase: return
        try:
            # We assume manifest.json is in a public bucket 'app-releases'
            import requests # We might need to import requests if not already there
            # Since we are using supabase client, we can try to get public url
            res = self.supabase.storage.from_("app-releases").download("manifest.json")
            if res:
                manifest = json.loads(res.decode('utf-8'))
                remote_version = manifest.get("version", "1.0.0")
                if remote_version > APP_VERSION:
                    update_url = manifest.get("url", "")
                    def notify():
                        msg = f"A new version ({remote_version}) is available!\n\n"
                        if update_url:
                            msg += f"Download here: {update_url}\n\n"
                        msg += "Would you like to visit the download page?"
                        if messagebox.askyesno("Update Available", msg):
                             import webbrowser
                             webbrowser.open(update_url or "https://github.com/PDF-Translator-Pro")
                    self.schedule_ui_update(notify)
        except Exception as e:
            print(f"Update Check Fail: {e}")

    def show_about(self):
        messagebox.showinfo("About", f"Bilingual Book Translator Pro\nVersion {APP_VERSION}\n\nA professional tool for PDF translation and dual-language book creation.\nCreated with Antigravity.")

    def toggle_settings_popdown(self):
        """Toggles the visibility of the settings popdown frame."""
        if self.settings_popdown and self.settings_popdown.winfo_exists():
            self.settings_popdown.destroy()
            self.settings_popdown = None
            return
            
        # Create Popdown Frame
        self.settings_popdown = ctk.CTkFrame(self, fg_color=CARD_BG, corner_radius=10, border_width=1, border_color=BORDER_COLOR, width=420)
        self.settings_popdown.place(relx=0.96, rely=0.07, anchor="ne")
        self.settings_popdown.lift()
        
        title_lbl = ctk.CTkLabel(self.settings_popdown, text="Translation Options", font=("Inter", 14, "bold"), text_color=TEXT_COLOR)
        title_lbl.pack(pady=(15, 5))
        
        info_var = ctk.StringVar(value="Hover over an option for details.")
        
        def on_enter(text):
            info_var.set(text)
            
        def on_leave(e):
            info_var.set("Hover over an option for details.")
            
        # 1-1 Pro Mode
        cb_one_one = ctk.CTkCheckBox(self.settings_popdown, text="1-1 Pro Mode (Bilingual Layout)", variable=self.one_one_var, command=self.toggle_one_one, text_color=TEXT_COLOR)
        cb_one_one.pack(anchor="w", padx=20, pady=10)
        cb_one_one.bind("<Enter>", lambda e: on_enter("Creates a strict side-by-side translation.\nOriginal left, translated right."))
        cb_one_one.bind("<Leave>", on_leave)

        # Deep Cleanup
        cb_deep = ctk.CTkCheckBox(self.settings_popdown, text="Deep Cleanup (Clean Mode)", variable=self.deep_cleanup_var, text_color=TEXT_COLOR)
        cb_deep.pack(anchor="w", padx=20, pady=10)
        cb_deep.bind("<Enter>", lambda e: on_enter("Removes headers/watermarks for cleaner output."))
        cb_deep.bind("<Leave>", on_leave)
        
        # Fast Mode
        cb_fast = ctk.CTkCheckBox(self.settings_popdown, text="Fast Mode (Low Quality OCR)", variable=self.speed_var, text_color=TEXT_COLOR)
        cb_fast.pack(anchor="w", padx=20, pady=10)
        cb_fast.bind("<Enter>", lambda e: on_enter("Speeds up processing by sacrificing layout retention."))
        cb_fast.bind("<Leave>", on_leave)
        
        # Eco Mode
        cb_eco = ctk.CTkCheckBox(self.settings_popdown, text="Eco Mode (Save Resources)", variable=self.eco_var, text_color=TEXT_COLOR)
        cb_eco.pack(anchor="w", padx=20, pady=10)
        cb_eco.bind("<Enter>", lambda e: on_enter("Reduces thread speed/memory for low-end devices."))
        cb_eco.bind("<Leave>", on_leave)

        # SPRINT 39: AI Engine Selection
        ctk.CTkLabel(self.settings_popdown, text="AI Translation Engine:", font=("Inter", 12, "bold")).pack(anchor="w", padx=20, pady=(10, 0))
        
        def on_engine_change(choice):
            self.selected_engine = choice
            self.save_settings()
            
        engine_menu = ctk.CTkOptionMenu(self.settings_popdown, values=["Google", "DeepL", "GPT-4o"], command=on_engine_change)
        engine_menu.set(self.selected_engine)
        engine_menu.pack(fill="x", padx=20, pady=5)
        engine_menu.bind("<Enter>", lambda e: on_enter("Choose your translation provider.\nGoogle is free, DeepL/GPT-4o require keys."))
        engine_menu.bind("<Leave>", on_leave)
        
        # DeepL Key Entry
        deepl_entry = ctk.CTkEntry(self.settings_popdown, placeholder_text="DeepL API Key...", height=30)
        deepl_entry.insert(0, self.deepl_key)
        deepl_entry.pack(fill="x", padx=20, pady=5)
        deepl_entry.bind("<KeyRelease>", lambda e: (setattr(self, 'deepl_key', deepl_entry.get()), self.save_settings()))
        deepl_entry.bind("<Enter>", lambda e: on_enter("Enter your DeepL API key (Free or Pro)."))
        deepl_entry.bind("<Leave>", on_leave)
        
        # OpenAI Key Entry
        openai_entry = ctk.CTkEntry(self.settings_popdown, placeholder_text="OpenAI API Key...", height=30)
        openai_entry.insert(0, self.openai_key)
        openai_entry.pack(fill="x", padx=20, pady=(5, 15))
        openai_entry.bind("<KeyRelease>", lambda e: (setattr(self, 'openai_key', openai_entry.get()), self.save_settings()))
        openai_entry.bind("<Enter>", lambda e: on_enter("Enter your OpenAI API key for GPT-4o."))
        openai_entry.bind("<Leave>", on_leave)

        # Glossary Manager Button
        ctk.CTkLabel(self.settings_popdown, text="Custom Terminology:", font=("Inter", 12, "bold")).pack(anchor="w", padx=20, pady=(5, 0))
        gloss_btn = ctk.CTkButton(self.settings_popdown, text="📖 Open Glossary Manager", 
                                 fg_color=ACCENT_COLOR, hover_color=ACCENT_HOVER, text_color="#FFFFFF",
                                 command=self.open_glossary_manager)
        gloss_btn.pack(fill="x", padx=20, pady=(5, 10))
        gloss_btn.bind("<Enter>", lambda e: on_enter("Set custom translations for names, places,\nor technical terms to keep them consistent."))
        gloss_btn.bind("<Leave>", on_leave)

        info_frame = ctk.CTkFrame(self.settings_popdown, fg_color=INPUT_BG, height=40, corner_radius=5)
        info_frame.pack(fill="x", padx=15, pady=(10, 15))
        info_frame.pack_propagate(False)

        info_lbl = ctk.CTkLabel(info_frame, textvariable=info_var, font=("Inter", 11, "italic"), text_color=TEXT_DIM)
        info_lbl.pack(expand=True)

    def create_menu_bar(self):
        """Create a professional Windows-style top menu bar."""
        self.menu_bar = tk.Menu(self)
        self.config(menu=self.menu_bar)

        # 1. File Menu
        self.file_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="File", menu=self.file_menu)
        self.file_menu.add_command(label="↗ Upload PDF...", command=self.upload_pdf, accelerator="Ctrl+O")
        self.file_menu.add_command(label="↺ Resume Last Session", command=self.resume_session)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="📚 Open Global Library", command=lambda: self.tab_view.set("Global Library"))
        self.file_menu.add_command(label="📥 Inbox (Requests)", command=self.show_inbox)
        self.file_menu.add_separator()
        self.file_menu.add_command(label="Exit", command=self.on_closing)

        # 2. View Menu (Toggles)
        self.view_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="View", menu=self.view_menu)
        self.view_menu.add_command(label="Toggle Theme (Light/Dark)", command=self.toggle_theme)
        self.view_menu.add_separator()
        self.view_menu.add_command(label="⚙ Preferences / Fonts", command=self.open_settings)

        # 3. Export Menu
        self.export_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Export", menu=self.export_menu)
        self.export_menu.add_command(label="Side-by-Side PDF", command=lambda: self.handle_export("Export PDF"))
        self.export_menu.add_command(label="Microsoft Word (.docx)", command=lambda: self.handle_export("Export Word"))
        self.export_menu.add_command(label="eBook (ePub)", command=lambda: self.handle_export("Export ePub"))

        # 4. Help Menu
        self.help_menu = tk.Menu(self.menu_bar, tearoff=0)
        self.menu_bar.add_cascade(label="Help", menu=self.help_menu)
        self.help_menu.add_command(label="Documentation", command=lambda: messagebox.showinfo("Help", "Visit documentation online for more details."))
        self.help_menu.add_separator()
        self.help_menu.add_command(label="About", command=self.show_about)

        self.bind_all("<Control-o>", lambda e: self.upload_pdf())

        self.is_paused = False
        self.stop_requested = False
        self.current_pdf_path = None
        self.bookmarks = [] # List of page indices
        self.current_theme = "Dark" # State tracker for theme toggle
        self.is_admin = False # SPRINT 34 Admin functionality

        # Storage and Pagination
        self.all_page_data = [] # List of {page, original, english, literal, ...}
        self.current_page_idx = 0
        self.total_pages = 0
        self.cache_dir = get_app_path(".translator_cache")
        if not os.path.exists(self.cache_dir):
            os.makedirs(self.cache_dir)
            
        # Cloud Sync Queue (prevents RAM spikes from spawning 100s of threads)
        self.sync_queue = []
        self.sync_worker_active = False
        self.sync_lock = threading.Lock()
            
        # LRU Image Cache (max 5 images to prevent RAM crashes)
        self.image_cache = {}
            
        # UI Recycling References
        self.page_number_label = ctk.CTkLabel(self, text="")
        self.orig_text = ctk.CTkTextbox(self)
        self.lit_text = ctk.CTkTextbox(self)
        self.visual_img_label = ctk.CTkLabel(self, text="")
        self.trans_text = ctk.CTkTextbox(self)
        
        # SPRINT 33: Workspace Tabs
        self.tab_view = ctk.CTkTabview(self, fg_color=PAPER_BG)
        self.tab_view.grid(row=1, column=0, rowspan=2, sticky="nsew", padx=10, pady=5)
        
        self.tab_pdf = self.tab_view.add("PDF Translator")
        self.tab_quick = self.tab_view.add("Quick Translate")
        self.tab_library = self.tab_view.add("Global Library")
        
        # Configure PDF Tab layout
        self.tab_pdf.grid_columnconfigure(0, weight=1)
        self.tab_pdf.grid_rowconfigure(1, weight=1)

        # Initialize Library UI
        self.build_library_ui()
        
        self.orig_container = ctk.CTkFrame(self.tab_pdf)
        self.mid_container = ctk.CTkFrame(self.tab_pdf)
        self.trans_container = ctk.CTkFrame(self.tab_pdf)

        # Headers Frame (Sticky) inside PDF Tab
        self.header_frame = ctk.CTkFrame(self.tab_pdf, fg_color=PAPER_BG, corner_radius=0)
        self.header_frame.grid(row=0, column=0, sticky="ew", padx=20)
        
        # Main Container (PDF Tab)
        self.main_container = ctk.CTkFrame(self.tab_pdf, fg_color=PAPER_BG, corner_radius=0)
        self.main_container.grid(row=1, column=0, sticky="nsew", padx=2, pady=2)
        
        # Debounce timer for text edits
        self.save_timer = None
        
        # Build the permanent UI skeleton once
        self.build_ui_skeleton()
        self.build_quick_translate_ui()
        
        # Footer Frame for Navigation inside PDF Tab
        self.footer_frame = ctk.CTkFrame(self.tab_pdf, fg_color=HEADER_BG, height=50, corner_radius=0)
        self.footer_frame.grid(row=2, column=0, sticky="ew")
        self.tab_pdf.grid_rowconfigure(2, weight=0)
        
        self.nav_first = ctk.CTkButton(self.footer_frame, text="⇤ First", width=80, fg_color=BTN_SECONDARY, text_color=TEXT_COLOR, hover_color=BTN_SECONDARY_HOVER, command=lambda: self.change_page("first"))
        self.nav_first.pack(side="left", padx=10, pady=10)
        
        self.nav_prev = ctk.CTkButton(self.footer_frame, text="← Prev", width=80, fg_color=BTN_SECONDARY, text_color=TEXT_COLOR, hover_color=BTN_SECONDARY_HOVER, command=lambda: self.change_page("prev"))
        self.nav_prev.pack(side="left", padx=5, pady=10)
        
        self.page_label = ctk.CTkLabel(self.footer_frame, text="Page 0 of 0", text_color=TEXT_COLOR, font=("Inter", 12, "bold"))
        self.page_label.pack(side="left", padx=20, pady=10)
        
        self.nav_next = ctk.CTkButton(self.footer_frame, text="Next →", width=80, fg_color=BTN_SECONDARY, text_color=TEXT_COLOR, hover_color=BTN_SECONDARY_HOVER, command=lambda: self.change_page("next"))
        self.nav_next.pack(side="left", padx=5, pady=10)
        
        self.nav_last = ctk.CTkButton(self.footer_frame, text="Last ⇥", width=80, fg_color=BTN_SECONDARY, text_color=TEXT_COLOR, hover_color=BTN_SECONDARY_HOVER, command=lambda: self.change_page("last"))
        self.nav_last.pack(side="left", padx=10, pady=10)

        self.jump_label = ctk.CTkLabel(self.footer_frame, text="Jump:", text_color=TEXT_DIM)
        self.jump_label.pack(side="left", padx=(50, 5), pady=10)
        
        self.jump_entry = ctk.CTkEntry(self.footer_frame, width=50, fg_color=INPUT_BG, border_color=BORDER_COLOR, text_color=TEXT_COLOR)
        self.jump_entry.pack(side="left", padx=5, pady=10)
        self.jump_entry.bind("<Return>", lambda e: self.change_page("jump"))

        self.rescan_err_btn = ctk.CTkButton(self.footer_frame, text="⚠ Fix Errors", width=100, 
                                            fg_color=BTN_DANGER, hover_color=BTN_DANGER_HOVER, text_color="#FFFFFF", command=self.rescan_all_errors)
        self.rescan_err_btn.pack(side="right", padx=10, pady=10)

        self.rescan_btn = ctk.CTkButton(self.footer_frame, text="↻ Rescan Page", width=100, 
                                        fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER, text_color=TEXT_COLOR, command=self.rescan_current_page)
        self.rescan_btn.pack(side="right", padx=10, pady=10)

        self.bookmark_btn = ctk.CTkButton(self.footer_frame, text="★ Bookmark", width=100,
                                         fg_color=CARD_BG, text_color=ACCENT_COLOR, hover_color=BTN_SECONDARY_HOVER,
                                         command=self.toggle_bookmark)
        self.bookmark_btn.pack(side="right", padx=10, pady=10)

        self.update_column_config()
        self.setup_headers()
        self.row_counter = 1
        
        # Check for existing session at startup
        self.check_initial_cache()
        
        # Trigger Login UI if Supabase is connected but no user
        if self.supabase and not self.current_user:
            self.after(500, self.update_avatar_ui)

    # --- SPRINT 34 & 35: Auth Methods & Profile ---
    def handle_avatar_click(self):
        if not self.current_user:
            self.show_auth_window()
        else:
            self.show_account_dash()
            
    def update_avatar_ui(self):
        """Updates the top right avatar button with the user's PFP or default icon."""
        if self.current_user:
            username = self.current_username or "User"
            if getattr(self, "avatar_url", None):
                # Load small 32x32 image
                img = self.get_cached_avatar(self.current_user["id"], self.avatar_url, size=(32, 32))
                if img:
                    self.avatar_btn.configure(image=img, text=username, compound="top")
                else:
                    self.avatar_btn.configure(image="", text=f"👤\n{username}")
            else:
                self.avatar_btn.configure(image="", text=f"👤\n{username}")
        else:
            self.avatar_btn.configure(image="", text="👤 Login")
            
    def show_account_dash(self):
        if not self.supabase or not self.current_user: return
        
        dash = ctk.CTkToplevel(self)
        dash.title("Account Dashboard")
        dash.geometry("400x500")
        dash.transient(self)
        dash.grab_set()
        dash.configure(fg_color=PAPER_BG)

        # Header
        name = self.current_username or "User"
        header = ctk.CTkLabel(dash, text=f"Welcome, {name}!", font=("Inter", 20, "bold"), text_color=TEXT_COLOR)
        header.pack(pady=(30, 10))

        # Avatar Display
        self.dash_avatar_lbl = ctk.CTkLabel(dash, text="👤", font=("Inter", 60), text_color=ACCENT_COLOR, fg_color=CARD_BG, corner_radius=10, width=120, height=120)
        self.dash_avatar_lbl.pack(pady=10)
        
        if getattr(self, "avatar_url", None):
            img = self.get_cached_avatar(self.current_user["id"], self.avatar_url, size=(100, 100))
            if img:
                self.dash_avatar_lbl.configure(image=img, text="")
            else:
                self.dash_avatar_lbl.configure(text="[Avatar Set]", font=("Inter", 14))

        # Update Avatar btn
        update_btn = ctk.CTkButton(dash, text="Update Profile Picture", command=self.upload_avatar, fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER, text_color=TEXT_COLOR)
        update_btn.pack(pady=(10, 30))

        # Inbox Button
        inbox_btn = ctk.CTkButton(dash, text="📬 Access Inbox (Requests)", command=self.show_inbox, height=40, font=("Inter", 14, "bold"))
        inbox_btn.pack(fill="x", padx=40, pady=10)

        # Logout Button
        logout_btn = ctk.CTkButton(dash, text="Logout", command=lambda: self.logout(dash), fg_color=BTN_DANGER, hover_color=BTN_DANGER_HOVER, text_color="#FFFFFF")
        logout_btn.pack(side="bottom", pady=20)

    def upload_avatar(self):
        file_path = filedialog.askopenfilename(title="Select Profile Picture", filetypes=[("Image Files", "*.png *.jpg *.jpeg")])
        if not file_path: return
        
        try:
            import time
            from PIL import Image
            
            # Compress image before upload
            img = Image.open(file_path)
            img.thumbnail((256, 256)) # resize to max 256x256
            
            import io
            img_byte_arr = io.BytesIO()
            img.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            file_name = f"{self.current_user['id']}_{int(time.time())}.png"
            
            # Upload to Supabase Storage
            self.supabase.storage.from_("avatars").upload(file_name, img_bytes, {"content-type": "image/png"})
            
            # Get public URL
            public_url = self.supabase.storage.from_("avatars").get_public_url(file_name)
            
            # Update DB
            self.supabase.table("profiles").update({"avatar_url": public_url}).eq("id", self.current_user["id"]).execute()
            
            self.avatar_url = public_url
            
            # Immediately refresh UI
            img = self.get_cached_avatar(self.current_user["id"], public_url, size=(100, 100), force_refresh=True)
            if img: self.dash_avatar_lbl.configure(image=img, text="")
            self.update_avatar_ui()
            
            messagebox.showinfo("Success", "Profile picture updated!")
            
        except Exception as e:
            messagebox.showerror("Upload Error", f"Failed to upload avatar: {e}")

    def logout(self, window):
        self.current_user = None
        self.current_username = None
        self.avatar_url = None
        self.is_admin = False
        if os.path.exists("session.json"):
            os.remove("session.json")
        self.update_avatar_ui()
        self.refresh_library()
        window.destroy()

    # --- SPRINT 34: Auth Methods ---
    def load_auth_session(self) -> Optional[Dict[str, str]]:
        path = get_app_path("session.json")
        if os.path.exists(path):
            try:
                with open(path, "r") as f:
                    return json.load(f)
            except: pass
        return None
        
    def save_auth_session(self, session) -> None:
        if not session: return
        try:
            path = get_app_path("session.json")
            with open(path, "w") as f:
                json.dump({"access_token": session.access_token, "refresh_token": session.refresh_token}, f)
        except Exception as e:
            print(f"Failed to save session: {e}")
            
    def fetch_username(self):
        if not self.supabase or not self.current_user: return
        try:
            # Join with profiles to get username
            res = self.supabase.table("profiles").select("*").eq("id", self.current_user["id"]).execute()
            if res.data:
                prof = res.data[0]
                self.current_username = prof.get('username')
                self.is_admin = prof.get('is_admin', False)
                self.avatar_url = prof.get('avatar_url')
        except Exception as e:
            print(f"Failed to fetch username: {e}")

    def show_auth_window(self):
        """Displays a modal Login/Register window."""
        auth_win = ctk.CTkToplevel(self)
        auth_win.title("Welcome to Translator Pro")
        auth_win.geometry("400x550")
        auth_win.grab_set()
        auth_win.attributes('-topmost', True)
        
        ctk.CTkLabel(auth_win, text="Translator Pro Cloud", font=("Inter", 24, "bold"), text_color=ACCENT_COLOR).pack(pady=(40, 10))
        ctk.CTkLabel(auth_win, text="Login to access the Global Library", font=("Inter", 12)).pack(pady=(0, 10))
        
        self.auth_error_lbl = ctk.CTkLabel(auth_win, text="", text_color="#E74C3C", font=("Inter", 12, "bold"))
        self.auth_error_lbl.pack(pady=(0, 10))
        
        tab_view = ctk.CTkTabview(auth_win, width=320)
        tab_view.pack(padx=20, pady=5, fill="both", expand=True)
        
        tab_login = tab_view.add("Login")
        tab_reg = tab_view.add("Register")
        
        # --- LOGIN TAB ---
        ctk.CTkLabel(tab_login, text="Email or Username:").pack(anchor="w", padx=20, pady=(10, 0))
        log_email = ctk.CTkEntry(tab_login, width=280)
        log_email.pack(padx=20, pady=5)
        
        ctk.CTkLabel(tab_login, text="Password:").pack(anchor="w", padx=20, pady=(10, 0))
        log_pass = ctk.CTkEntry(tab_login, width=280, show="*")
        log_pass.pack(padx=20, pady=5)
        
        def _do_login():
            login_id = log_email.get().strip()
            pw = log_pass.get().strip()
            if not login_id or not pw: return
            
            log_btn.configure(state="disabled", text="Logging in...")
            self.auth_error_lbl.configure(text="")
            try:
                em = login_id
                if "@" not in login_id:
                    # Attempt to resolve username to email via profiles table
                    user_res = self.supabase.table("profiles").select("email").eq("username", login_id).execute()
                    if user_res.data and user_res.data[0].get("email"):
                        em = user_res.data[0]["email"]
                    else:
                        raise Exception("Username not found or has no linked email. Please login with Email.")
                        
                res = self.supabase.auth.sign_in_with_password({"email": em, "password": pw})
                if res and res.user and res.session:
                    self.current_user = {"id": res.user.id, "email": res.user.email}
                    
                    # Backfill email if missing in profiles for future username logins
                    try:
                        self.supabase.table("profiles").update({"email": res.user.email}).eq("id", res.user.id).execute()
                    except: pass
                    
                    self.save_auth_session(res.session)
                    self.fetch_username()
                    self.update_avatar_ui()
                    auth_win.destroy()
                    self.refresh_library() # Refresh to show permissions
            except Exception as e:
                err_msg = str(e)
                if "Invalid login" in err_msg: err_msg = "Invalid email or password."
                self.auth_error_lbl.configure(text=err_msg)
            finally:
                log_btn.configure(state="normal", text="Login")
                
        log_btn = ctk.CTkButton(tab_login, text="Login", command=_do_login, width=280)
        log_btn.pack(pady=30)
        
        # --- REGISTER TAB ---
        ctk.CTkLabel(tab_reg, text="Username:").pack(anchor="w", padx=20, pady=(10, 0))
        reg_user = ctk.CTkEntry(tab_reg, width=280)
        reg_user.pack(padx=20, pady=5)
        
        ctk.CTkLabel(tab_reg, text="Email:").pack(anchor="w", padx=20, pady=(10, 0))
        reg_email = ctk.CTkEntry(tab_reg, width=280)
        reg_email.pack(padx=20, pady=5)
        
        ctk.CTkLabel(tab_reg, text="Password:").pack(anchor="w", padx=20, pady=(10, 0))
        reg_pass = ctk.CTkEntry(tab_reg, width=280, show="*")
        reg_pass.pack(padx=20, pady=5)
        
        def _do_register():
            usr = reg_user.get().strip()
            em = reg_email.get().strip()
            pw = reg_pass.get().strip()
            
            if not usr or not em or len(pw) < 6:
                self.auth_error_lbl.configure(text="Please fill all fields. Password must be 6+ chars.")
                return
                
            reg_btn.configure(state="disabled", text="Registering...")
            self.auth_error_lbl.configure(text="")
            try:
                # 1. Sign up
                res = self.supabase.auth.sign_up({"email": em, "password": pw})
                if res and res.user:
                    user_id = res.user.id
                    # 2. Create Profile
                    try:
                        self.supabase.table("profiles").insert({"id": user_id, "username": usr, "email": em}).execute()
                    except Exception as prof_err:
                        print(f"Profile creation err: {prof_err}")
                        
                    # 3. Direct Login
                    login_res = self.supabase.auth.sign_in_with_password({"email": em, "password": pw})
                    if login_res and login_res.user and login_res.session:
                        self.current_user = {"id": user_id, "email": em}
                        self.current_username = usr
                        self.save_auth_session(login_res.session)
                        self.update_avatar_ui()
                        auth_win.destroy()
                        self.refresh_library()
            except Exception as e:
                self.auth_error_lbl.configure(text=str(e))
            finally:
                reg_btn.configure(state="normal", text="Create Account")
                
        reg_btn = ctk.CTkButton(tab_reg, text="Create Account", command=_do_register, width=280, fg_color="#4CAF50", hover_color="#45A049")
        reg_btn.pack(pady=20)
        
        # --- Skip Option ---
        ctk.CTkButton(auth_win, text="Skip for now (Local Mode Only)", fg_color="transparent", 
                      hover_color="#333333", command=auth_win.destroy).pack(pady=10)

    def toggle_theme(self):
        """Switches the application between Light and Dark mode."""
        if self.current_theme == "Dark":
            ctk.set_appearance_mode("Light")
            self.current_theme = "Light"
        else:
            ctk.set_appearance_mode("Dark")
            self.current_theme = "Dark"

    def load_settings(self):
        self.font_family = "Times New Roman"
        self.font_size = 15
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, "r", encoding="utf-8") as f:
                    config = json.load(f)
                    self.font_family = config.get("font_family", "Times New Roman")
                    self.font_size = config.get("font_size", 15)
                    self.bookmarks = config.get("bookmarks", [])
                    self.selected_engine = config.get("selected_engine", "Google")
                    self.deepl_key = config.get("deepl_key", "")
                    self.openai_key = config.get("openai_key", "")
                    self.glossary = config.get("glossary", {})
            except:
                pass

    def save_settings(self):
        config = {
            "font_family": self.font_family,
            "font_size": self.font_size,
            "bookmarks": self.bookmarks,
            "selected_engine": self.selected_engine,
            "deepl_key": self.deepl_key,
            "openai_key": self.openai_key,
            "glossary": self.glossary
        }
        try:
            with open(self.config_file, "w", encoding="utf-8") as f:
                json.dump(config, f)
        except:
            pass
            
    def search_text(self, event=None):
        query = self.search_var.get().lower()
        if not query or not self.all_page_data: return
        
        # Start searching from the next page
        start_idx = (self.current_page_idx + 1) % len(self.all_page_data)
        for i in range(len(self.all_page_data)):
            idx = (start_idx + i) % len(self.all_page_data)
            data = self.all_page_data[idx]
            text_pool = (str(data.get('original', '')) + 
                         str(data.get('english', '')) + 
                         str(data.get('literal', ''))).lower()
            if query in text_pool:
                self.current_page_idx = idx
                self.render_page(idx)
                self.status_label.configure(text=f"Found '{query}' on Page {idx+1}")
                return
        
        messagebox.showinfo("Search", f"No more matches found for '{query}'.")

    def open_glossary_manager(self):
        """Opens a window to manage custom terminology."""
        gloss_win = ctk.CTkToplevel(self)
        gloss_win.title("Terminology Glossary Manager")
        gloss_win.geometry("500x600")
        gloss_win.attributes("-topmost", True)
        gloss_win.configure(fg_color=CARD_BG)
        
        ctk.CTkLabel(gloss_win, text="Custom Glossary", font=("Inter", 18, "bold"), text_color=TEXT_COLOR).pack(pady=20)
        ctk.CTkLabel(gloss_win, text="Define rules like: 'Geralt' -> 'Geralt of Rivia'\nEnsures consistent translation across all engines.", 
                    font=("Inter", 12), text_color=TEXT_DIM).pack(pady=(0, 20))
        
        # List Frame
        list_frame = ctk.CTkScrollableFrame(gloss_win, fg_color=INPUT_BG, height=350)
        list_frame.pack(fill="both", padx=20, expand=True)
        
        def refresh_list():
            for widget in list_frame.winfo_children():
                widget.destroy()
            
            for orig, target in self.glossary.items():
                row = ctk.CTkFrame(list_frame, fg_color="transparent")
                row.pack(fill="x", pady=2)
                ctk.CTkLabel(row, text=f"{orig} → {target}", font=("Arial", 13)).pack(side="left", padx=10)
                
                def _del(o=orig):
                    del self.glossary[o]
                    self.save_settings()
                    refresh_list()
                    
                ctk.CTkButton(row, text="✕", width=30, height=25, fg_color="#E53935", 
                             command=_del).pack(side="right", padx=5)

        # Add Frame
        add_frame = ctk.CTkFrame(gloss_win, fg_color="transparent")
        add_frame.pack(fill="x", padx=20, pady=20)
        
        orig_entry = ctk.CTkEntry(add_frame, placeholder_text="Original Word/Name...", width=160)
        orig_entry.pack(side="left", padx=5)
        
        target_entry = ctk.CTkEntry(add_frame, placeholder_text="Target Translation...", width=160)
        target_entry.pack(side="left", padx=5)
        
        def _add():
            o = orig_entry.get().strip()
            t = target_entry.get().strip()
            if o and t:
                self.glossary[o] = t
                self.save_settings()
                orig_entry.delete(0, 'end')
                target_entry.delete(0, 'end')
                refresh_list()
            else:
                messagebox.showwarning("Input Error", "Please provide both words.")
                
        ctk.CTkButton(add_frame, text="+ Add", width=80, fg_color=ACCENT_COLOR, 
                     command=_add).pack(side="left", padx=5)
        
        refresh_list()

    def toggle_bookmark(self):
        if self.current_page_idx in self.bookmarks:
            self.bookmarks.remove(self.current_page_idx)
            self.status_label.configure(text=f"Removed Bookmark for Page {self.current_page_idx + 1}")
        else:
            self.bookmarks.append(self.current_page_idx)
            self.bookmarks.sort()
            self.status_label.configure(text=f"Bookmarked Page {self.current_page_idx + 1}")
        self.update_bookmark_ui()
        self.save_settings()

    def update_bookmark_ui(self):
        if self.current_page_idx in self.bookmarks:
            self.bookmark_btn.configure(text="★ Bookmarked", fg_color="#f39c12")
        else:
            self.bookmark_btn.configure(text="☆ Bookmark", fg_color="#f1c40f")

    def build_library_ui(self):
        """Builds the nested UI for the Global Library tab."""
        # Main container
        container = ctk.CTkFrame(self.tab_library, fg_color="transparent")
        container.pack(fill="both", expand=True, padx=20, pady=20)

        # Header with Search and Refresh
        header = ctk.CTkFrame(container, fg_color="transparent")
        header.pack(fill="x", pady=(0, 20))

        self.lib_search_var = ctk.StringVar()
        search_entry = ctk.CTkEntry(header, placeholder_text="Search Library...", 
                                    textvariable=self.lib_search_var, width=250)
        search_entry.pack(side="left", padx=(0, 10))
        search_entry.bind("<KeyRelease>", lambda e: self.refresh_library())
        
        self.lib_filter_var = ctk.StringVar(value="Public Library")
        filter_seg = ctk.CTkSegmentedButton(header, values=["My Books", "Public Library"], variable=self.lib_filter_var, command=lambda e: self.refresh_library())
        filter_seg.pack(side="left", padx=10)

        refresh_btn = ctk.CTkButton(header, text="↻ Refresh", width=100, command=self.refresh_library)
        refresh_btn.pack(side="right")

        # Scrollable area for books
        self.lib_scroll = ctk.CTkScrollableFrame(container, fg_color=PAPER_BG)
        self.lib_scroll.pack(fill="both", expand=True)

        # Initial Load
        self.refresh_library()

    def refresh_library(self):
        """SPRINT 33/40: Fetches books from Supabase with account-based sync and permission checks."""
        if not self.supabase: return

        # Clear existing
        for widget in self.lib_scroll.winfo_children():
            widget.destroy()

        search_term = self.lib_search_var.get().lower()

        try:
            # Query Logic
            show_mode = self.lib_filter_var.get() if hasattr(self, "lib_filter_var") else "Public Library"
            
            if show_mode == "My Books" and getattr(self, "current_user", None):
                # ACCOUNT SYNC: Pull everything owned by THIS account
                res = self.supabase.table("books").select("*").eq("owner_id", self.current_user["id"]).order("created_at", desc=True).execute()
            else:
                # PUBLIC LIBRARY: Pull everything marked public
                res = self.supabase.table("books").select("*").eq("is_public", True).order("created_at", desc=True).execute()
            
            books = res.data
            if not books:
                ctk.CTkLabel(self.lib_scroll, text="No books found in this view.").pack(pady=40)
                return
                
            local_cache_dir = get_app_path(".translator_cache")
            local_caches = set(os.listdir(local_cache_dir)) if os.path.exists(local_cache_dir) else set()

            # Group by Category
            raw_folders: Dict[str, List[Any]] = {}
            for book in books:
                if search_term and search_term not in book['title'].lower() and search_term not in book.get('language', '').lower():
                    continue
                
                cat_key = str(book.get('category', 'Uncategorized') or 'Uncategorized')
                if cat_key not in raw_folders: raw_folders[cat_key] = []
                raw_folders[cat_key].append(book)

            folders = cast(Dict[str, List[Any]], raw_folders)
            if not folders:
                ctk.CTkLabel(self.lib_scroll, text="No matches found.").pack(pady=40)
                return

            # Render Folders (Explorer Style Grid)
            for category in sorted(folders.keys()):
                cat_frame = ctk.CTkFrame(self.lib_scroll, fg_color=FRAME_BG, corner_radius=10)
                cat_frame.pack(pady=10, fill="x", padx=10)

                cat_header = ctk.CTkLabel(cat_frame, text=f"📂 {category}", 
                                          font=("Inter", 16, "bold"), text_color=("#6200EE", "#BB86FC"), anchor="w")
                cat_header.pack(fill="x", padx=15, pady=10)

                book_grid = ctk.CTkFrame(cat_frame, fg_color="transparent")
                book_grid.pack(fill="x", padx=10, pady=(0, 10))

                for i, book in enumerate(folders[category]):
                    card = ctk.CTkFrame(book_grid, fg_color=CARD_BG, width=280, height=120)
                    card.grid(row=i // 3, column=i % 3, padx=10, pady=10, sticky="nsew")
                    book_grid.grid_columnconfigure(i % 3, weight=1)

                    is_owner = self.current_user and book.get('owner_id') == self.current_user['id']
                    is_locked = book.get('is_read_only', True)
                    
                    # Icons
                    status_icon = "👤" if is_owner else "🔓" if not is_locked else "🔒"

                    # Info
                    info_frame = ctk.CTkFrame(card, fg_color="transparent")
                    info_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)

                    ctk.CTkLabel(info_frame, text=book['title'], font=("Inter", 13, "bold"), anchor="w", wraplength=180).pack(fill="x")
                    
                    uploader = book.get('owner_username') or "Anonymous"
                    sub_text = f"{book['language']} • {book.get('total_pages','?')} Pgs\nBy: {uploader}"
                    ctk.CTkLabel(info_frame, text=sub_text, font=("Inter", 11), text_color=TEXT_DIM, anchor="w", justify="left").pack(fill="x")

                    # Actions Frame
                    act_frame = ctk.CTkFrame(card, fg_color="transparent")
                    act_frame.pack(side="right", padx=10)

                    # 1. READ/EDIT Button
                    if not is_locked or is_owner:
                        btn_txt = "Edit" if not is_locked or is_owner else "Read"
                        ctk.CTkButton(act_frame, text=btn_txt, width=70, height=28,
                                      fg_color=BTN_PRIMARY, hover_color=BTN_PRIMARY_HOVER,
                                      command=lambda b=book: self.load_book_from_cloud(b)).pack(pady=2)
                    else:
                        # 2. REQUEST ACCESS Button (Collaborative Workflow)
                        ctk.CTkButton(act_frame, text="Request Edit", width=70, height=28,
                                      fg_color=BTN_WARNING, hover_color=BTN_WARNING_HOVER, text_color="#121212",
                                      command=lambda b=book: self.request_edit_access(b['id'], b['title'], b['owner_id'])).pack(pady=2)
                        
                        ctk.CTkButton(act_frame, text="Read Only", width=70, height=28,
                                      fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER, text_color=TEXT_COLOR,
                                      command=lambda b=book: self.load_book_from_cloud(b)).pack(pady=2)

                    if getattr(self, "is_admin", False) or is_owner:
                        ctk.CTkButton(act_frame, text="🗑", width=30, height=28, fg_color=BTN_DANGER, 
                                      command=lambda b=book['id']: self.delete_cloud_book(b, self)).pack(pady=2)

        except Exception as e:
            ctk.CTkLabel(self.lib_scroll, text=f"Connection Error: {e}", text_color="#E74C3C").pack(pady=40)

        except Exception as e:
            ctk.CTkLabel(self.lib_scroll, text=f"Error connecting to Cloud: {e}", text_color="#E74C3C").pack(pady=40)

    def open_library(self):
        """Legacy method for popup - now switches to tab."""
        self.tab_view.set("Global Library")
        self.refresh_library()

    def get_cached_avatar(self, user_id, avatar_url, size=(32, 32), force_refresh=False):
        import urllib.request
        from PIL import Image, ImageDraw
        cache_dir = os.path.join(os.getcwd(), ".translator_cache", "avatars")
        os.makedirs(cache_dir, exist_ok=True)
        file_path = os.path.join(cache_dir, f"{user_id}_{size[0]}.png")
        
        if not os.path.exists(file_path) or force_refresh:
            try:
                # Need absolute reliable cache invalidation if force refreshing
                req = urllib.request.Request(avatar_url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req) as response, open(file_path, 'wb') as out_file:
                    out_file.write(response.read())
                    
                # Make it round
                img = Image.open(file_path).convert("RGBA")
                img = img.resize(size, Image.Resampling.LANCZOS)
                mask = Image.new("L", size, 0)
                draw = ImageDraw.Draw(mask)
                draw.ellipse((0, 0, size[0], size[1]), fill=255)
                img.putalpha(mask)
                img.save(file_path, "PNG")
            except Exception as e:
                print(f"Avatar fetch failed: {e}")
                return None
                
        try:
            return ctk.CTkImage(light_image=Image.open(file_path), dark_image=Image.open(file_path), size=size)
        except:
            return None

    def delete_cloud_book(self, book_id, window):
        """Allows user to delete their own uploaded books."""
        if not self.supabase: return
        if not messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this book from the cloud?"): return
        
        try:
            self.supabase.table("books").delete().eq("id", book_id).execute()
            messagebox.showinfo("Deleted", "Book removed from cloud.")
            self.refresh_library()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete: {e}")

    def build_quick_translate_ui(self):
        """SPRINT 33: Builds the side-by-side Google Translate style layout for the Quick Translate tab."""
        self.tab_quick.grid_columnconfigure(0, weight=1)
        self.tab_quick.grid_columnconfigure(1, weight=1)
        self.tab_quick.grid_rowconfigure(1, weight=1)
        
        # --- Top Controls ---
        top_bar = ctk.CTkFrame(self.tab_quick, fg_color="transparent")
        top_bar.grid(row=0, column=0, columnspan=2, sticky="ew", pady=(0, 10))
        
        # Left (Source Language)
        left_ctrl = ctk.CTkFrame(top_bar, fg_color="transparent")
        left_ctrl.pack(side="left", fill="x", expand=True)
        ctk.CTkLabel(left_ctrl, text="Translate from:", font=("Inter", 12, "bold")).pack(side="left", padx=10)
        self.qt_src_lang = ctk.CTkComboBox(left_ctrl, values=["Auto-Detect", "Arabic", "Hebrew", "French", "English", "Greek", "Latin"], width=150)
        self.qt_src_lang.set("Auto-Detect")
        self.qt_src_lang.pack(side="left")
        
        # Right (Target Language)
        right_ctrl = ctk.CTkFrame(top_bar, fg_color="transparent")
        right_ctrl.pack(side="right", fill="x", expand=True)
        ctk.CTkLabel(right_ctrl, text="Translate to:", font=("Inter", 12, "bold")).pack(side="left", padx=10)
        self.qt_tgt_lang = ctk.CTkComboBox(right_ctrl, values=["English", "Arabic", "French", "Spanish"], width=150)
        self.qt_tgt_lang.set("English")
        self.qt_tgt_lang.pack(side="left")
        
        # --- Text Areas ---
        # Input Box
        self.qt_input = ctk.CTkTextbox(self.tab_quick, font=(self.font_family, self.font_size))
        self.qt_input.grid(row=1, column=0, sticky="nsew", padx=(0, 5))
        self.qt_input.insert("1.0", "Paste or type text here to translate...")
        
        # Output Box
        self.qt_output = ctk.CTkTextbox(self.tab_quick, font=(self.font_family, self.font_size), fg_color=INPUT_BG)
        self.qt_output.grid(row=1, column=1, sticky="nsew", padx=(5, 0))
        self.qt_output.configure(state="disabled")
        
        # --- Bottom Action Bar ---
        bottom_bar = ctk.CTkFrame(self.tab_quick, fg_color="transparent")
        bottom_bar.grid(row=2, column=0, columnspan=2, sticky="ew", pady=(10, 0))
        
        self.qt_btn = ctk.CTkButton(bottom_bar, text="Translate Now", font=("Inter", 14, "bold"), 
                                    height=40, width=200, fg_color="#1976D2", hover_color="#1565C0",
                                    command=self.run_quick_translate)
        self.qt_btn.pack(pady=10)

    def run_quick_translate(self):
        """Executes translation for the Quick Translate tab."""
        raw_text = self.qt_input.get("1.0", "end-1c").strip()
        if not raw_text or raw_text == "Paste or type text here to translate...":
            return
            
        src_lang = self.qt_src_lang.get()
        tgt_lang = self.qt_tgt_lang.get()
        
        # Map nice names to deep_translator codes
        code_map = {
            "Auto-Detect": "auto", "Arabic": "ar", "Hebrew": "he", "French": "fr", 
            "English": "en", "Greek": "el", "Latin": "la", "Spanish": "es"
        }
        
        src_code = code_map.get(src_lang, "auto")
        tgt_code = code_map.get(tgt_lang, "en")
        
        self.qt_btn.configure(state="disabled", text="Translating...")
        self.qt_output.configure(state="normal")
        self.qt_output.delete("1.0", "end")
        self.qt_output.insert("1.0", "Translating...")
        self.qt_output.configure(state="disabled")
        
        def _translate_worker():
            try:
                # 1. Formatting pre-process (Numeral Shield etc)
                protected_text = raw_text
                if src_code in ["ar", "he"] or src_code == "auto":
                    # Simple heuristic: if text looks RTL, apply shield
                    if len(re.findall(r'[\u0600-\u06FF\u0590-\u05FF]', protected_text)) > 5:
                        protected_text = self.normalize_to_western_digits(protected_text)
                        
                        # Apply Bidi Isolation to numbers
                        def freeze_number(match):
                            return f"\u2066{match.group(0)}\u2069"
                        protected_text = re.sub(r'\b[\d\.\-\:]+\b', freeze_number, protected_text)
                
                # 2. Translate
                translated_text = ""
                if self.selected_engine == "DeepL" and self.deepl_key:
                    import deepl
                    t = deepl.Translator(self.deepl_key)
                    # DeepL uses 'auto' differently or uses source_lang
                    src = None if src_code == 'auto' else src_code.upper()
                    # Check target lang mapping for DeepL
                    target_map = {"en": "EN-US", "es": "ES", "fr": "FR", "de": "DE"}
                    tgt = target_map.get(tgt_code, "EN-US")
                    res = t.translate_text(protected_text, target_lang=tgt, source_lang=src)
                    translated_text = res.text
                elif self.selected_engine == "GPT-4o" and self.openai_key:
                    from openai import OpenAI
                    client = OpenAI(api_key=self.openai_key)
                    response = client.chat.completions.create(
                      model="gpt-4o",
                      messages=[
                        {"role": "system", "content": f"You are a professional translator. Translate the following text from {src_code} to {tgt_code}. Preserve formatting. Output ONLY the translated text."},
                        {"role": "user", "content": protected_text}
                      ]
                    )
                    translated_text = response.choices[0].message.content.strip()
                else: # Default Google
                    translator = GoogleTranslator(source=src_code, target=tgt_code)
                    if len(protected_text) > 4000:
                        chunks = [protected_text[i:i+4000] for i in range(0, len(protected_text), 4000)]
                        for chunk in chunks:
                            translated_text += translator.translate(chunk) + " "
                    else:
                        translated_text = translator.translate(protected_text)
                    
                # 3. Post-Process (remove bidi markers)
                if translated_text:
                     translated_text = translated_text.replace('\u2066', '').replace('\u2069', '')
                     
                self.schedule_ui_update(lambda: self._finalize_qt(translated_text))
            except Exception as e:
                self.schedule_ui_update(lambda: self._finalize_qt(f"Error: {e}"))
                
        threading.Thread(target=_translate_worker, daemon=True).start()
        
    def _finalize_qt(self, result_text):
        self.qt_output.configure(state="normal")
        self.qt_output.delete("1.0", "end")
        
        # Tag for RTL if output is Arabic/Hebrew
        if self.qt_tgt_lang.get() in ["Arabic", "Hebrew"]:
            self.qt_output.insert("1.0", result_text, "rtl")
            self.qt_output.tag_config("rtl", justify="right")
        else:
            self.qt_output.insert("1.0", result_text)
            
        self.qt_output.configure(state="disabled")
        self.qt_btn.configure(state="normal", text="Translate Now")

    def publish_to_gallery(self):
        """SPRINT 33/40: Opens the Explorer-style Publish Manager."""
        if not self.supabase or not self.current_pdf_path: return
        
        # 1. Create Popup Window
        pub_win = ctk.CTkToplevel(self)
        pub_win.title("Publish Manager")
        pub_win.geometry("450x550")
        pub_win.grab_set() 
        
        main_container = ctk.CTkFrame(pub_win, fg_color="transparent")
        main_container.pack(fill="both", expand=True, padx=25, pady=25)

        ctk.CTkLabel(main_container, text="Cloud Explorer", font=("Inter", 20, "bold")).pack(pady=(0, 20))
        
        # 2. Custom Title Input
        ctk.CTkLabel(main_container, text="Book Title:", font=("Inter", 12, "bold")).pack(anchor="w")
        title_var = ctk.StringVar(value=os.path.basename(self.current_pdf_path))
        ctk.CTkEntry(main_container, textvariable=title_var, width=400).pack(pady=(5, 15))
        
        # 3. Folder/Category Selection (Explorer Style)
        ctk.CTkLabel(main_container, text="Select Destination Folder:", font=("Inter", 12, "bold")).pack(anchor="w")
        
        folders_frame = ctk.CTkFrame(main_container, fg_color=FRAME_BG, height=150)
        folders_frame.pack(fill="x", pady=(5, 5))
        
        folders_list = ctk.CTkScrollableFrame(folders_frame, height=120, fg_color="transparent")
        folders_list.pack(fill="both", expand=True, padx=5, pady=5)

        # State for folder selection
        selected_folder = ctk.StringVar(value="Uncategorized")
        folder_btns = {}

        def _select_folder(name):
            selected_folder.set(name)
            for n, btn in folder_btns.items():
                if n == name:
                    btn.configure(fg_color=("#BB86FC", "#6200EE"), text_color="#121212")
                else:
                    btn.configure(fg_color="transparent", text_color=TEXT_COLOR)

        def _refresh_folders():
            for widget in folders_list.winfo_children(): widget.destroy()
            
            try:
                res = self.supabase.table("books").select("category").execute()
                names = set(["Uncategorized"])
                if res.data:
                    for b in res.data:
                        if b.get("category"): names.add(b["category"])
                
                for name in sorted(list(names)):
                    btn = ctk.CTkButton(folders_list, text=f"📂 {name}", anchor="w", 
                                        fg_color="transparent", text_color=TEXT_COLOR,
                                        hover_color=("#BB86FC", "#6200EE"), height=30,
                                        command=lambda n=name: _select_folder(n))
                    btn.pack(fill="x", pady=2)
                    folder_btns[name] = btn
                
                _select_folder(selected_folder.get())
            except: pass

        _refresh_folders()

        # Add New Folder Row
        new_folder_row = ctk.CTkFrame(main_container, fg_color="transparent")
        new_folder_row.pack(fill="x", pady=(0, 15))
        
        new_f_entry = ctk.CTkEntry(new_folder_row, placeholder_text="New folder name...", width=300)
        new_f_entry.pack(side="left")

        def _add_folder():
            name = new_f_entry.get().strip()
            if name:
                selected_folder.set(name)
                new_f_entry.delete(0, "end")
                _refresh_folders()

        add_btn = ctk.CTkButton(new_folder_row, text="+", width=40, font=("Inter", 16, "bold"),
                                fg_color=BTN_SECONDARY, hover_color=BTN_SECONDARY_HOVER,
                                command=_add_folder)
        add_btn.pack(side="right", padx=(5, 0))

        # 4. Privacy & Collaboration Toggles
        ctk.CTkLabel(main_container, text="Cloud Settings:", font=("Inter", 12, "bold")).pack(anchor="w")
        
        settings_box = ctk.CTkFrame(main_container, fg_color=FRAME_BG)
        settings_box.pack(fill="x", pady=(5, 20))

        # Public Release Toggle
        is_public_var = ctk.BooleanVar(value=False) # DEFAULT TO PRIVATE (Safe for Production)
        public_switch = ctk.CTkSwitch(settings_box, text="Release to Public Library", variable=is_public_var,
                                      progress_color=("#BB86FC", "#6200EE"))
        public_switch.pack(anchor="w", padx=15, pady=10)
        
        # Read-Only Toggle
        read_only_var = ctk.BooleanVar(value=True)
        read_only_switch = ctk.CTkSwitch(settings_box, text="Read-Only (Requires Edit Request)", variable=read_only_var)
        read_only_switch.pack(anchor="w", padx=15, pady=(0, 10))

        # 5. Submit Action
        def _execute_publish():
            pub_title = title_var.get().strip() or (os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Untitled")
            pub_folder = selected_folder.get().strip() or "Uncategorized"
            is_read_only = read_only_var.get()
            is_public = is_public_var.get()
            
            pub_btn.configure(state="disabled", text="Syncing to Cloud...")
            
            def _job():
                # Pass is_public to sync_book_metadata (need to update that method too)
                book_id_or_err = self.sync_book_metadata(pub_title, pub_folder, is_read_only, is_public)
                if not book_id_or_err or (isinstance(book_id_or_err, str) and book_id_or_err.startswith("Error:")):
                    err_hint = book_id_or_err if isinstance(book_id_or_err, str) else "Unknown"
                    self.schedule_ui_update(lambda msg=err_hint: messagebox.showerror("Cloud Error", f"Sync failed.\n\nReason: {msg}"))
                    self.schedule_ui_update(lambda: pub_btn.configure(state="normal", text="Sync Now"))
                    return
                book_id = book_id_or_err
                    
                try:
                    if self.current_pdf_path:
                        _book_name = os.path.basename(self.current_pdf_path).replace(".pdf", "")
                        self.cache_dir = get_app_path(os.path.join(".translator_cache", _book_name))
                    
                    self.schedule_ui_update(lambda: pub_btn.configure(text="Uploading Pages..."))
                    for page_data in self.all_page_data:
                        if "page" not in page_data: continue
                        page_idx = page_data["page"]
                        
                        self.supabase.table("pages").upsert({
                            "book_id": book_id,
                            "page_index": page_idx,
                            "original": page_data.get("original", ""),
                            "english": page_data.get("english", ""),
                            "literal": page_data.get("literal", ""),
                            "is_image": page_data.get("is_image", False),
                            "is_cover": page_data.get("is_cover", False),
                            "is_centered": page_data.get("is_centered", False),
                            "is_rtl_page": page_data.get("is_rtl_page", False)
                        }, on_conflict="book_id,page_index").execute()
                        
                        dyn_img_path = os.path.join(self.cache_dir, f"img_{page_idx}.png")
                        if os.path.exists(dyn_img_path):
                            try:
                                with open(dyn_img_path, "rb") as f:
                                    self.supabase.storage.from_("book-images").upload(
                                        f"{str(book_id)}/img_{int(page_idx)}.png",
                                        f, file_options={"content-type": "image/png", "upsert": "true"}
                                    )
                            except Exception: pass

                    self.schedule_ui_update(lambda: messagebox.showinfo("Success", f"Book synced to cloud folder '{pub_folder}'!"))
                    self.schedule_ui_update(lambda: self.status_label.configure(text="Cloud Sync Complete"))
                    self.schedule_ui_update(lambda: self.sync_btn.pack(side="right", padx=5, pady=10)) 
                    self.schedule_ui_update(lambda: self.sync_btn.configure(state="normal"))
                    self.schedule_ui_update(lambda: self.publish_btn.pack_forget())
                    self.schedule_ui_update(pub_win.destroy)
                    self.schedule_ui_update(self.refresh_library) # Refresh list
                except Exception as e:
                    self.schedule_ui_update(lambda err=e: messagebox.showerror("Error", f"Failed: {err}"))
                    self.schedule_ui_update(lambda: pub_btn.configure(state="normal", text="Sync Now"))
                    
            threading.Thread(target=_job, daemon=True).start()
            
        pub_btn = ctk.CTkButton(main_container, text="Sync to Cloud", height=40, font=("Inter", 14, "bold"),
                                fg_color=BTN_PRIMARY, hover_color=BTN_PRIMARY_HOVER, text_color="#FFFFFF", 
                                command=_execute_publish)
        pub_btn.pack(fill="x", pady=5)

    def force_cloud_sync(self):
        """Manually pushes local edits back to cloud."""
        if not self.supabase or not self.all_page_data or not self.current_pdf_path: return
        
        title = os.path.basename(self.current_pdf_path)
        self.sync_btn.configure(state="disabled", text="Syncing...")
        
        def _job():
            try:
                # sync_book_metadata now requires title
                book_id = self.sync_book_metadata(title)
                if not book_id:
                    self.schedule_ui_update(lambda: messagebox.showerror("Sync Failed", "Book not registered in cloud or sync failed."))
                    return
                
                # 1. Check permissions
                res = self.supabase.table("books").select("id, is_read_only, owner_id").eq("id", book_id).execute()
                if not res.data:
                    self.schedule_ui_update(lambda: messagebox.showerror("Sync Failed", "Book is not registered in the cloud."))
                    return
                    
                book_id = res.data[0]['id']
                is_read_only = res.data[0].get('is_read_only', True)
                
                if is_read_only:
                    # SPRINT 34: Check if they are the owner OR if they have an approved request
                    is_owner = self.current_user and res.data[0].get('owner_id') == self.current_user['id']
                    
                    if getattr(self, "is_admin", False):
                        is_owner = True # Override for admin
                        
                    if not is_owner:
                        # Check for approved request
                        allowed = False
                        if self.current_user:
                            try:
                                req_res = self.supabase.table("permission_requests").select("status").eq("book_id", book_id).eq("requester_id", self.current_user["id"]).eq("status", "approved").execute()
                                if req_res.data:
                                    allowed = True
                            except Exception as e:
                                print(f"Permission Check Error: {e}")
                                
                        if not allowed:
                            self.schedule_ui_update(lambda: messagebox.showwarning("Permission Denied", "This book is marked Read-Only. You must request edit access from the owner."))
                            return
                
                # 2. Push all cached edits to cloud
                for page_data in self.all_page_data:
                    if "page" not in page_data: continue
                    page_idx = page_data["page"]
                    
                    self.supabase.table("pages").upsert({
                        "book_id": book_id,
                        "page_index": page_idx,
                        "original": page_data.get("original", ""),
                        "english": page_data.get("english", ""),
                        "literal": page_data.get("literal", "")
                    }, on_conflict="book_id,page_index").execute()
                    
                self.schedule_ui_update(lambda: messagebox.showinfo("Sync Complete", "Collaborative edits saved to cloud!"))
            except Exception as e:
                self.schedule_ui_update(lambda err=e: messagebox.showerror("Sync Error", str(err)))
            finally:
                self.schedule_ui_update(lambda: self.sync_btn.configure(state="normal", text="⇪ Sync Edits"))
                
        threading.Thread(target=_job, daemon=True).start()

    def load_book_from_cloud(self, book_info):
        """Downloads book cache from cloud and replaces local cache asynchronously."""
        if not self.supabase: return
        
        self.status_label.configure(text=f"Pulling '{book_info['title']}' from cloud...")
        self.menu_bar.entryconfigure("Export", state="disabled")
        
        # Disable translation controls to prevent interference
        self.pause_btn.configure(state="disabled")
        self.rescan_btn.configure(state="disabled")
        
        def _download_worker():
            try:
                # 1. Clear local cache for this title
                book_id = book_info['id']
                # We assume current_pdf_path might not be set yet, but we'll use title to locate local cache
                # For simplicity, we'll just download all pages into a temp cache dir
                temp_cache = get_app_path(os.path.join(".translator_cache", book_info['title'].replace(".pdf", "")))
                os.makedirs(temp_cache, exist_ok=True)
                self.cache_dir = temp_cache
                
                # SPRINT 33: Also set the PDF path to match the cloud title so saving works
                self.current_pdf_path = os.path.join(os.getcwd(), book_info['title'])
                
                res = self.supabase.table("pages").select("*").eq("book_id", book_id).order("page_index").execute()
                new_page_data = []
                for p in res.data:
                    p_data = {
                        "page": p['page_index'],
                        "original": p['original'],
                        "english": p['english'],
                        "literal": p['literal'],
                        "is_image": p.get('is_image', False),
                        "is_cover": p.get('is_cover', False),
                        "is_centered": p.get('is_centered', False),
                        "is_rtl_page": p.get('is_rtl_page', False)
                    }
                    # Restore image download logic inside the background thread
                    img_path = os.path.join(temp_cache, f"img_{p['page_index']}.png")
                    if not os.path.exists(img_path):
                        try:
                            # Try to download the page's image
                            res_img = self.supabase.storage.from_("book-images").download(f"{book_id}/img_{p['page_index']}.png")
                            with open(img_path, 'wb') as f:
                                f.write(res_img)
                            p_data['cover_image_path'] = img_path
                        except Exception as e:
                            print(f"Debug: Missed image for page {p['page_index']} - {e}")
                            pass
                    else:
                        p_data['cover_image_path'] = img_path
                        
                    new_page_data.append(p_data)
                
                if not new_page_data:
                    self.schedule_ui_update(lambda: messagebox.showwarning("Empty Book", "This cloud book has no translated pages yet."))
                    self.schedule_ui_update(lambda: self.status_label.configure(text="Ready"))
                    return
                
                # Update UI elements back on the main thread
                def _finalize_load():
                    self.all_page_data = new_page_data
                    self.total_pages = len(self.all_page_data)
                    self.current_page_idx = 0
                    
                    # Force tabs back to PDF Translator
                    self.tab_view.set("PDF Translator")
                    
                    # Update buttons
                    self.sync_btn.pack(side="right", padx=5, pady=10) # Reveal Sync Button
                    self.sync_btn.configure(state="normal")
                    self.publish_btn.pack_forget() # Hide Publish for cloud books
                    
                    self.menu_bar.entryconfigure("Export", state="normal")
                    self.pause_btn.configure(state="normal")
                    
                    # SPRINT 34: Permissions & Request Flow
                    is_locked = book_info.get('is_read_only', True)
                    is_owner = self.current_user and book_info.get('owner_id') == self.current_user['id']
                    
                    if getattr(self, "is_admin", False):
                        is_owner = True # Override for admin
                        is_locked = False
                        
                    # If it's locked and we aren't the owner, show the request button
                    if is_locked and not is_owner:
                        self.sync_btn.configure(state="disabled") # Disable sync explicitly
                        self.publish_btn.configure(text="Request Edit Access", fg_color="#F39C12", hover_color="#D68910", text_color="#FFFFFF", state="normal")
                        
                        # Replace publish command with request command for this session
                        self.publish_btn.configure(command=lambda: self.request_edit_access(book_info['id'], book_info['title'], book_info.get('owner_id')))
                        self.publish_btn.pack(side="right", padx=10, pady=10)
                    elif is_owner and is_locked:
                        self.publish_btn.pack_forget() 
                        self.sync_btn.configure(state="normal") # Owner can sync their own locked book
                        # We don't unlock the UI boxes for the owner because `is_locked` is globally respected in `set_box_lock` below, unless we override it. Let's override for the owner.
                        is_locked = False 
                    else:
                        self.publish_btn.pack_forget() # Hide if unlocked or we own it (and it's not locked against us)
                        self.sync_btn.configure(state="normal")
                    
                    def set_box_lock(box, locked):
                        box.configure(state="normal")
                        if locked: box.configure(state="disabled")

                    set_box_lock(self.orig_text, is_locked)
                    set_box_lock(self.lit_text, is_locked)
                    set_box_lock(self.trans_text, is_locked)

                    self.render_page(0)
                    self.status_label.configure(text=f"Loaded '{book_info['title']}' from Cloud.")
                    
                self.schedule_ui_update(_finalize_load)
                
            except Exception as e:
                self.schedule_ui_update(lambda err=e: messagebox.showerror("Cloud Error", f"Failed to pull book: {err}"))
                self.schedule_ui_update(lambda: self.status_label.configure(text="Ready"))
                
        # Start background thread
        threading.Thread(target=_download_worker, daemon=True).start()

    def request_edit_access(self, book_id, book_title, owner_id):
        """Sends a permission request to the book owner."""
        if not self.supabase or not self.current_user:
            messagebox.showwarning("Login Required", "You must be logged in to request access.")
            return
            
        if not owner_id:
            messagebox.showerror("Error", "This book has no assigned owner.")
            return
            
        # Check if already requested
        try:
            res = self.supabase.table("permission_requests").select("status").eq("book_id", book_id).eq("requester_id", self.current_user["id"]).execute()
            if res.data:
                status = res.data[0]['status']
                if status == 'pending':
                    messagebox.showinfo("Already Requested", "You have a pending request for this book.")
                    return
                elif status == 'approved':
                    messagebox.showinfo("Already Approved", "Your request is already approved. Please refresh the library.")
                    return
        except Exception as e:
            print(f"Check Request Error: {e}")
            
        # Insert new request
        try:
            self.supabase.table("permission_requests").insert({
                "book_id": book_id,
                "owner_id": owner_id,
                "requester_id": self.current_user["id"],
                "requester_username": self.current_username,
                "book_title": book_title
            }).execute()
            messagebox.showinfo("Request Sent", f"Edit access requested for '{book_title}'.\nThe owner will be notified.")
            self.publish_btn.configure(text="Request Pending", state="disabled")
        except Exception as e:
            messagebox.showerror("Request Failed", f"Failed to send request: {e}")

    def show_inbox(self):
        """SPRINT 34: Shows incoming permission requests."""
        if not self.supabase or not self.current_user: return
        
        inbx = ctk.CTkToplevel(self)
        inbx.title("Publisher Inbox")
        inbx.geometry("500x400")
        inbx.transient(self)
        inbx.grab_set()
        inbx.configure(fg_color=PAPER_BG)
        
        ctk.CTkLabel(inbx, text="Incoming Edit Requests", font=("Inter", 18, "bold"), text_color=TEXT_COLOR).pack(pady=(20, 10))
        
        scroll = ctk.CTkScrollableFrame(inbx, width=450, height=300, fg_color="transparent")
        scroll.pack(pady=10, padx=20, fill="both", expand=True)
        
        try:
            # Get books owned by this user
            books_res = self.supabase.table("books").select("id, title").eq("owner_id", self.current_user["id"]).execute()
            if not books_res.data:
                ctk.CTkLabel(scroll, text="You have not published any books.", text_color=TEXT_DIM).pack(pady=20)
                return
                
            book_map = {b['id']: b['title'] for b in books_res.data}
            book_ids = list(book_map.keys())
            
            # Use 'in' filter to get requests for these books
            reqs = self.supabase.table("permission_requests").select("*").in_("book_id", book_ids).eq("status", "pending").execute()
            
            if not reqs.data:
                ctk.CTkLabel(scroll, text="No pending requests.", text_color=TEXT_DIM).pack(pady=20)
                return
                
            for req in reqs.data:
                # Build Request Card...
                req_card = ctk.CTkFrame(scroll, fg_color=CARD_BG, corner_radius=8)
                req_card.pack(fill="x", pady=5)
                
                info = ctk.CTkFrame(req_card, fg_color="transparent")
                info.pack(side="left", padx=10, pady=10)
                
                b_title = book_map.get(req['book_id'], "Unknown Book")
                ctk.CTkLabel(info, text=f"From: {req['requester_username']}", font=("Inter", 12, "bold"), text_color=TEXT_COLOR).pack(anchor="w")
                ctk.CTkLabel(info, text=f"Book: {b_title}", font=("Inter", 11), text_color=TEXT_DIM).pack(anchor="w")
                
                acts = ctk.CTkFrame(req_card, fg_color="transparent")
                acts.pack(side="right", padx=10)
                
                def _approve(r_id=req['id']):
                    try:
                        self.supabase.table("permission_requests").update({"status": "approved"}).eq("id", r_id).execute()
                        inbx.destroy()
                        self.show_inbox() # Refresh
                        messagebox.showinfo("Success", "Request approved.")
                    except Exception as e:
                        messagebox.showerror("Error", str(e))
                        
                def _deny(r_id=req['id']):
                    try:
                        self.supabase.table("permission_requests").update({"status": "denied"}).eq("id", r_id).execute()
                        inbx.destroy()
                        self.show_inbox() # Refresh
                    except Exception as e:
                        messagebox.showerror("Error", str(e))
                        
                ctk.CTkButton(acts, text="Approve", width=60, fg_color=BTN_SECONDARY, hover_color="#4CAF50", command=_approve).pack(side="left", padx=5)
                ctk.CTkButton(acts, text="Deny", width=60, fg_color=BTN_DANGER, hover_color=BTN_DANGER_HOVER, command=_deny).pack(side="left")
                
        except Exception as e:
            err_msg = str(e)
            print(f"Inbox Fetch Error: {err_msg}")
            # Gracefully degrade to empty state for any database fetch errors to avoid scaring the user
            ctk.CTkLabel(scroll, text="No pending requests.", text_color=TEXT_DIM).pack(pady=40)

    # --- SPRINT 27: Intelligent Language Detection ---
    def detect_script(self, text):
        """Returns the likely Tesseract OCR code(s) joined by '+' and RTL status."""
        if not text or len(text) < 5:
            return {"ocr": "eng", "trans": "auto", "rtl": False}
            
        # Character ranges with broader detection
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
        hebrew_chars = len(re.findall(r'[\u0590-\u05FF]', text))
        greek_chars = len(re.findall(r'[\u0370-\u03FF]', text))
        cyrillic_chars = len(re.findall(r'[\u0400-\u04FF]', text))
        chinese_chars = len(re.findall(r'[\u4E00-\u9FFF]', text))
        latin_chars = len(re.findall(r'[a-zA-Z\u00C0-\u017F]', text))
        
        counts = {
            "ara": arabic_chars,
            "heb": hebrew_chars,
            "ell": greek_chars,
            "rus": cyrillic_chars,
            "chi_sim": chinese_chars,
            "eng": latin_chars
        }
        
        # Mapping back to translator codes for single-language fallbacks
        trans_map = {
            "ara": "ar", "heb": "he", "ell": "el", "rus": "ru", "chi_sim": "zh-CN", "eng": "en", "fra": "fr"
        }
        
        total_chars = sum(counts.values())
        if total_chars == 0:
            return {"ocr": "eng", "trans": "auto", "rtl": False}
            
        # Identify all significant scripts (> 2% of characters or > 2 chars total)
        detected_ocr_codes = []
        is_rtl = False
        
        # Sort to keep eng/latin/fra at the end (Tesseract convention)
        potential_codes = ["ara", "heb", "ell", "rus", "chi_sim", "fra", "eng"]
        
        # Priority Hack: If any latin is found, use 'fra' instead of 'eng' as it handles accents better
        if counts["eng"] > 0:
            counts["fra"] = counts["eng"]
            counts["eng"] = 0

        for code in potential_codes:
            char_count = counts.get(code, 0)
            if char_count > 2 or (char_count > 0 and char_count / total_chars > 0.02):
                if code == "fra":
                    # Always include eng with fra for safety
                    detected_ocr_codes.extend(["fra", "eng"])
                else:
                    detected_ocr_codes.append(code)
                
                if code in ["ara", "heb"]:
                    is_rtl = True
        
        if not detected_ocr_codes:
            detected_ocr_codes = ["eng"]
            
        final_ocr = "+".join(detected_ocr_codes)
        
        # Determine likely primary source language for the translator
        # If multiple, use 'auto'
        if len(detected_ocr_codes) > 1:
            final_trans = "auto"
        else:
            final_trans = trans_map.get(detected_ocr_codes[0], "auto")
            
        return {
            "ocr": final_ocr,
            "trans": final_trans,
            "rtl": is_rtl
        }

    def open_settings(self):
        settings_win = ctk.CTkToplevel(self)
        settings_win.title("Text Settings")
        settings_win.geometry("300x250")
        settings_win.attributes("-topmost", True)
        
        # Font Family
        ctk.CTkLabel(settings_win, text="Font Family:").pack(pady=(20, 5))
        fonts = ["Times New Roman", "Helvetica", "Arial", "Courier New", "Georgia", "Verdana"]
        font_menu = ctk.CTkOptionMenu(settings_win, values=fonts, command=self.update_font_family)
        font_menu.set(self.font_family)
        font_menu.pack(pady=5)
        
        # Font Size
        ctk.CTkLabel(settings_win, text="Font Size:").pack(pady=(10, 5))
        size_slider = ctk.CTkSlider(settings_win, from_=10, to=30, number_of_steps=20, command=self.update_font_size)
        size_slider.set(self.font_size)
        size_slider.pack(pady=5)

        # Bookmarks List
        if self.bookmarks:
            ctk.CTkLabel(settings_win, text="Jump to Bookmark:").pack(pady=(10, 5))
            bm_vals = [f"Page {i+1}" for i in self.bookmarks]
            bm_menu = ctk.CTkOptionMenu(settings_win, values=bm_vals, 
                                        command=lambda v: [self.goto_bookmark(v), settings_win.destroy()])
            bm_menu.set("Select Bookmark")
            bm_menu.pack(pady=5)

    def goto_bookmark(self, val):
        page_num = int(val.split(" ")[1]) - 1
        self.current_page_idx = page_num
        self.render_page(page_num)

    def update_font_family(self, choice):
        self.font_family = choice
        self.apply_font_settings()
        
    def update_font_size(self, value):
        self.font_size = int(value)
        self.apply_font_settings()
        
    def apply_font_settings(self):
        font_tpl = (self.font_family, self.font_size)
        lit_font_tpl = (self.font_family, max(10, self.font_size - 1)) # Lit is slightly smaller
        
        self.orig_text.configure(font=font_tpl)
        self.trans_text.configure(font=font_tpl)
        self.lit_text.configure(font=lit_font_tpl)
        self.save_settings()

    def check_initial_cache(self):
        """Checks if there's an existing cache on disk and shows Resume button if so."""
        import glob
        base_cache = os.path.join(os.getcwd(), ".translator_cache")
        
        # Disable resume initially
        self.file_menu.entryconfigure("↺ Resume Last Session", state="disabled")
        
        if os.path.exists(base_cache):
            # Search all book subfolders for cached pages
            for sub in os.listdir(base_cache):
                sub_path = os.path.join(base_cache, sub)
                if os.path.isdir(sub_path):
                    files = glob.glob(os.path.join(sub_path, "page_*.json"))
                    if files:
                        self.cache_dir = sub_path  # Point to the correct book subfolder
                        self.file_menu.entryconfigure("↺ Resume Last Session", state="normal")
                        self.status_label.configure(text=f"Found existing session ({len(files)} pages).")
                        return

    def resume_session(self):
        """Loads all page JSON files from cache directly into memory."""
        import glob
        import json
        import re
        
        files = glob.glob(os.path.join(self.cache_dir, "page_*.json"))
        if not files: return
        
        self.status_label.configure(text="Loading session from disk...")
        self.all_page_data = []
        self.image_cache = {}
        
        def load_cache():
            try:
                # Sort files by the page number in their filename
                def extract_num(f):
                    match = re.search(r'page_(\d+)\.json', f)
                    return int(match.group(1)) if match else 0
                    
                sorted_files = sorted(files, key=extract_num)
                
                for f in sorted_files:
                    with open(f, 'r', encoding='utf-8') as jfile:
                        data = json.load(jfile)
                        page_idx = data.get('page', 0)
                        
                        # Ensure the list is large enough to insert at this index
                        while len(self.all_page_data) <= page_idx:
                            self.all_page_data.append(None)
                            
                        self.all_page_data[page_idx] = data
                        
                self.total_pages = len(self.all_page_data)
                
                # Ready to view
                self.schedule_ui_update(lambda: self.status_label.configure(text="Session Restored!"))
                self.schedule_ui_update(lambda: self.menu_bar.entryconfigure("Export", state="normal"))
                
                # Set the app to a paused state so the user can click "Resume"
                self.is_paused = True
                self.stop_requested = True
                self.schedule_ui_update(lambda: self.pause_btn.configure(text="Resume", fg_color="#2E7D32", hover_color="#1B5E20", command=self.toggle_pause))
                
                # Automatically pull the original PDF path from the first page's metadata
                if self.all_page_data and 'original_pdf_path' in self.all_page_data[0]:
                    cache_path = self.all_page_data[0]['original_pdf_path']
                    if os.path.exists(cache_path):
                        self.current_pdf_path = cache_path
                    else:
                        # Fallback if they moved the file
                        self.schedule_ui_update(lambda: messagebox.showwarning("File Moved", f"Could not find original PDF at:\n{cache_path}\nRe-scan feature will be disabled."))
                
                if self.all_page_data:
                    self.current_page_idx = 0
                    self.schedule_ui_update(lambda: self.render_page(self.current_page_idx))
                    self.schedule_ui_update(lambda: self.page_label.configure(text=f"Page 1 of {self.total_pages}"))
                    
            except Exception as e:
                self.schedule_ui_update(lambda msg=str(e): messagebox.showerror("Resume Error", msg))
            finally:
                self.schedule_ui_update(lambda: self.status_label.configure(text="Ready"))
                
        threading.Thread(target=load_cache, daemon=True).start()

    def toggle_one_one(self):
        self.update_column_config()
        self.setup_headers()
        # Rerender current page to update layout (Visual Feed vs 1-1 Match)
        if self.all_page_data:
            self.render_page(self.current_page_idx)

    def update_column_config(self):
        # Sync weights between header and container
        # Always maintain 3 columns + index
        for area in [self.header_frame, self.main_container]:
            area.grid_columnconfigure(0, weight=1) # Page #
            area.grid_columnconfigure(1, weight=5) # Transcription
            area.grid_columnconfigure(2, weight=5) # 1-1 OR Visual Feed
            area.grid_columnconfigure(3, weight=5) # English
        
    def setup_headers(self):
        # Clear existing headers first to avoid stacking
        for widget in self.header_frame.winfo_children():
            widget.destroy()
            
        h_font = ("Inter", 14, "bold")
        ctk.CTkLabel(self.header_frame, text="#", font=h_font, text_color=ACCENT_COLOR).grid(row=0, column=0, pady=10, sticky="n")
        ctk.CTkLabel(self.header_frame, text="Transcription", font=h_font, text_color=TEXT_COLOR).grid(row=0, column=1, pady=10)
        
        # Dynamic Middle Header
        mid_text = "1-1 Match Pro" if self.one_one_var.get() else "Source Page"
        ctk.CTkLabel(self.header_frame, text=mid_text, font=h_font, text_color=ACCENT_COLOR).grid(row=0, column=2, pady=10)
            
        ctk.CTkLabel(self.header_frame, text="English Translation", font=h_font, text_color=TEXT_COLOR).grid(row=0, column=3, pady=10)
        
    def upload_pdf(self):
        file_path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if file_path:
            self.status_label.configure(text="Processing...")
            
            # Clear previous entries visually from the recycled ui
            self.page_number_label.configure(text="Processing new document...")
            
            if self.supabase:
                self.publish_btn.configure(state="normal")
                self.sync_btn.configure(state="normal")
                
            def clear_box(box):
                box.configure(state="normal")
                box.delete("1.0", "end")
                box.configure(state="disabled")
                
            clear_box(self.orig_text)
            clear_box(self.lit_text)
            clear_box(self.trans_text)
            self.visual_img_label.configure(image="", text="[Loading...]")
            self.stop_requested = True
            if self.active_executor:
                try: self.active_executor.shutdown(wait=False, cancel_futures=True)
                except: pass
            if self.ui_pdf_doc:
                try: self.ui_pdf_doc.close()
                except: pass
                self.ui_pdf_doc = None
            
            self.current_session += 1
            self.current_pdf_path = file_path
            self.stop_requested = False
            
            # Reset state for new book
            self.all_page_data = [] 
            self.image_cache = {} 
            self.current_page_idx = 0
            self.total_pages = 0            
            # Start everything in background to keep UI fluid
            threading.Thread(target=self.initial_upload_kickoff, args=(file_path,), daemon=True).start()

    def initial_upload_kickoff(self, file_path):
        """Heavy background lifting before starting full PDF processing."""
        # 1. Cloud Check (Network call - must be background)
        if self.supabase:
            book_title = os.path.basename(file_path)
            try:
                res = self.supabase.table("books").select("*").eq("title", book_title).execute()
                if res.data:
                    def ask_cloud():
                        use_cloud = messagebox.askyesno("Cloud Match Found", 
                                                        f"'{book_title}' has existing progress in the cloud. \n\nDo you want to pull cloud translations instead of starting fresh?")
                        if use_cloud:
                            self.load_book_from_cloud(res.data[0])
                        else:
                            threading.Thread(target=self.start_fresh_processing, args=(file_path,), daemon=True).start()
                    self.schedule_ui_update(ask_cloud)
                    return
            except: pass

        self.start_fresh_processing(file_path)

    def start_fresh_processing(self, file_path):
        """Clears local cache and starts OCR pipeline."""
        book_name = os.path.basename(file_path).replace(".pdf", "")
        self.cache_dir = get_app_path(os.path.join(".translator_cache", book_name))
        os.makedirs(self.cache_dir, exist_ok=True)
        
        # 2. Cache Cleanup (Disk I/O)
        import glob
        for f in glob.glob(os.path.join(self.cache_dir, "*")):
            try: 
                if os.path.isfile(f): os.remove(f)
            except: pass
            
        # 3. Start OCR
        self.process_pdf(file_path)

    def change_page(self, direction):
        if not self.all_page_data:
            return
            
        # FORCE IMMEDIATE SAVE OF CURRENT PAGE BEFORE NAVIGATING AWAY
        self.save_edits_to_cache()
        
        new_idx = self.current_page_idx
        if direction == "first":
            new_idx = 0
        elif direction == "last":
            new_idx = len(self.all_page_data) - 1
        elif direction == "next":
            if self.current_page_idx < len(self.all_page_data) - 1:
                new_idx += 1
        elif direction == "prev":
            if self.current_page_idx > 0:
                new_idx -= 1
        elif direction == "jump":
            try:
                page_num = int(self.jump_entry.get())
                if 1 <= page_num <= len(self.all_page_data):
                    new_idx = page_num - 1
            except:
                pass
                
        if new_idx != self.current_page_idx or direction == "jump":
            print(f"[Nav] Moving to Index {new_idx} (Page {new_idx+1})")
            self.current_page_idx = int(new_idx)
            self.render_page(self.current_page_idx)

    def save_page_to_cache(self, page_data):
        # Save images separately if they exist to keep RAM lean
        page_num = page_data['page']
        cache_path = os.path.join(self.cache_dir, f"page_{page_num}.json")
        
        # Strip PIL images from JSON, save them as PNGs
        save_data = page_data.copy()
        if save_data.get('cover_image'):
            img_path = os.path.join(self.cache_dir, f"img_{page_num}.png")
            save_data['cover_image'].save(img_path)
            save_data['cover_image_path'] = img_path
            # We explicitly delete it from the original page_data dictionary in memory!
            # The render_page function will lazy load it back if needed
            del save_data['cover_image']
            page_data['cover_image_path'] = img_path
            
        # Add the original PDF path to the cache
        if self.current_pdf_path:
            save_data['original_pdf_path'] = self.current_pdf_path
            page_data['original_pdf_path'] = self.current_pdf_path
            
        import json
        with open(cache_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)
            
            
        # Add to Sync Queue (sequential to save RAM)
        if self.supabase:
            with self.sync_lock:
                self.sync_queue.append(save_data)
                if not self.sync_worker_active:
                    self.sync_worker_active = True
                    threading.Thread(target=self.process_sync_queue, daemon=True).start()

    def process_sync_queue(self):
        """Worker thread that handles cloud sync sequentially."""
        while True:
            item = None
            with self.sync_lock:
                if self.sync_queue:
                    item = self.sync_queue.pop(0)
                else:
                    self.sync_worker_active = False
                    break
            
            if item:
                try:
                    self.sync_page_to_cloud(item)
                except:
                    pass
        
        # Explicit cleanup after batch sync
        import gc
        gc.collect()

    def sync_book_metadata(self, book_title, category="Uncategorized", is_read_only=True, is_public=False):
        """Registers or updates book information in Supabase."""
        if not self.supabase or not self.current_user:
            return "Error: Not logged in"
        
        try:
            # Check for existing
            res = self.supabase.table("books").select("id").eq("title", book_title).eq("owner_id", self.current_user["id"]).execute()
            if res.data:
                book_id = res.data[0]["id"]
                self.supabase.table("books").update({
                    "category": category,
                    "is_read_only": is_read_only,
                    "is_public": is_public
                }).eq("id", book_id).execute()
                return book_id
                
            # Register new book
            new_book = {
                "title": book_title,
                "total_pages": len(self.all_page_data),
                "language": self.lang_menu.get(),
                "category": category,
                "is_read_only": is_read_only,
                "is_public": is_public,
                "owner_id": self.current_user["id"],
                "owner_username": self.current_username
            }
            res = self.supabase.table("books").insert(new_book).execute()
            return res.data[0]['id']
        except Exception as e:
            error_msg = str(e)
            print(f"Metadata Sync Error: {error_msg}")
            return f"Error: {error_msg}"

    def sync_page_to_cloud(self, page_data):
        """Uploads page JSON and any images to Supabase."""
        if not self.supabase: return
        
        book_id = self.sync_book_metadata()
        if not book_id: return
        
        page_num = page_data['page']
        try:
            # 1. Sync DB
            payload = {
                "book_id": book_id,
                "page_index": page_num,
                "original": page_data['original'],
                "english": page_data['english'],
                "literal": page_data['literal'],
                "is_image": page_data['is_image'],
                "is_cover": page_data['is_cover'],
                "is_centered": page_data.get('is_centered', False),
                "is_rtl_page": page_data.get('is_rtl_page', False)
            }
            
            # Upsert: Update if book_id+page_index exists
            self.supabase.table("pages").upsert(payload, on_conflict="book_id,page_index").execute()
            
            # 2. Sync Image to Storage
            dyn_img_path = os.path.join(self.cache_dir, f"img_{page_num}.png")
            if os.path.exists(dyn_img_path):
                file_name = f"{book_id}/img_{page_num}.png"
                try:
                    with open(dyn_img_path, 'rb') as f:
                        self.supabase.storage.from_("book-images").upload(
                            file_name, f, file_options={"content-type": "image/png", "upsert": "true"}
                        )
                except Exception as img_e:
                    print(f"Failed to upload image for page {page_num}: {img_e}")
        except Exception as e:
            print(f"Cloud Sync Error (Page {page_num}): {e}")


    def build_ui_skeleton(self):
        # Clear main container
        for widget in self.main_container.winfo_children():
            widget.destroy()

        self.main_container.grid_columnconfigure(0, weight=1) # Page#
        self.main_container.grid_columnconfigure(1, weight=5) # Transcription
        self.main_container.grid_columnconfigure(2, weight=5) # 1-1 OR Visual Feed
        self.main_container.grid_columnconfigure(3, weight=5) # English
        self.main_container.grid_rowconfigure(1, weight=1) # Allow text boxes to expand vertically
        
        # Column 0: Page Number
        self.page_number_label = ctk.CTkLabel(self.main_container, text="Page 0", font=("Inter", 13, "bold"), text_color=ACCENT_COLOR)
        self.page_number_label.grid(row=1, column=0, padx=10, pady=20, sticky="n")

        # Create Paper Sheet Containers for Column 1, 2, 3
        def create_paper_sheet(col):
            frame = ctk.CTkFrame(self.main_container, fg_color=PAPER_SHEET_BG, corner_radius=12, border_width=1, border_color=BORDER_COLOR)
            frame.grid(row=1, column=col, padx=8, pady=15, sticky="nsew")
            
            # Use a slightly darker color for the textboxes to create depth
            box = ctk.CTkTextbox(frame, fg_color="transparent", text_color=TEXT_COLOR, 
                                 font=("Inter", 14), wrap="word", border_width=0, corner_radius=0)
            box.pack(padx=15, pady=15, fill="both", expand=True)
            return box, frame

        self.orig_text, self.orig_container = create_paper_sheet(1)
        self.lit_text, self.mid_container = create_paper_sheet(2)
        self.trans_text, self.trans_container = create_paper_sheet(3)
        
        # Bind key release events for all textboxes for auto-save
        self.orig_text.bind("<KeyRelease>", lambda e: self.schedule_save_edits())
        self.lit_text.bind("<KeyRelease>", lambda e: self.schedule_save_edits())
        self.trans_text.bind("<KeyRelease>", lambda e: self.schedule_save_edits())
        
        # Special Label for image pages in column 2 (Visual Feed)
        self.visual_img_label = ctk.CTkLabel(self.mid_container, text="", fg_color="transparent")
        # Hidden initially
        self.visual_img_label.pack_forget()
        
    def schedule_save_edits(self):
        """Sets a debounced timer so we don't save to disk on every single keystroke."""
        if self.save_timer is not None:
            self.after_cancel(self.save_timer)
        self.save_timer = self.after(1500, self.save_edits_to_cache)
        
    def save_edits_to_cache(self):
        """Pulls the current text from all 3 boxes and rewrites the cache JSON."""
        if not self.all_page_data or self.current_page_idx >= len(self.all_page_data):
            return
            
        # SPRINT 33 Fix: If the page is currently processing (None), don't try to save edits to it
        if self.all_page_data[self.current_page_idx] is None:
            return
            
        # Get live text (ignoring the trailing newline Tkinter always adds)
        new_orig = self.orig_text.get("1.0", "end-1c")
        new_lit = self.lit_text.get("1.0", "end-1c")
        new_trans = self.trans_text.get("1.0", "end-1c")
        
        # Update memory struct
        self.all_page_data[self.current_page_idx]['original'] = new_orig
        self.all_page_data[self.current_page_idx]['literal'] = new_lit
        self.all_page_data[self.current_page_idx]['english'] = new_trans
        
        # Save straight to disk
        self.save_page_to_cache(self.all_page_data[self.current_page_idx])

    def render_page(self, page_idx):
        try:
            if not self.all_page_data or page_idx >= len(self.all_page_data): return
            data = self.all_page_data[page_idx]
            
            # Handle Placeholder (Not processed yet)
            if data is None:
                self.populate_ui_skeleton(
                    page_num=page_idx, 
                    original="[Page Still Processing...]", 
                    translated="Please wait for OCR to complete.", 
                    literal="[Processing...]", 
                    is_image=False, 
                    is_rtl_override=False, 
                    is_centered=True, 
                    is_cover=False, 
                    cover_image=None
                )
                self.page_label.configure(text=f"Page {page_idx + 1} of {len(self.all_page_data)}")
                return

            page_id = data['page']
            
            # Smart Memory Management (LRU Cache)
            img_ref = self.image_cache.get(page_id)
            
            if not img_ref and data.get('cover_image_path') and os.path.exists(data['cover_image_path']):
                try: img_ref = Image.open(data['cover_image_path'])
                except: pass

            # ULTIMATE FALLBACK: Extract directly from PDF if image is still missing
            if not img_ref and self.current_pdf_path:
                try:
                    if not self.ui_pdf_doc:
                        self.ui_pdf_doc = fitz.open(self.current_pdf_path)
                    page = self.ui_pdf_doc.load_page(page_idx)
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    img_ref = Image.open(io.BytesIO(pix.tobytes("png")))
                except Exception as e:
                    print(f"[Render Fallback Error] Page {page_idx}: {e}")
                    self.ui_pdf_doc = None 

            if img_ref:
                self.image_cache[page_id] = img_ref
                if len(self.image_cache) > 5:
                    oldest_key = next(iter(self.image_cache))
                    self.image_cache.pop(oldest_key, None)

            self.populate_ui_skeleton(
                page_num=data['page'], 
                original=data['original'], 
                translated=data['english'], 
                literal=data['literal'], 
                is_image=data['is_image'], 
                is_rtl_override=data['is_rtl_page'], 
                is_centered=data['is_centered'], 
                is_cover=data['is_cover'], 
                cover_image=img_ref
            )
            
            self.page_label.configure(text=f"Page {page_idx + 1} of {len(self.all_page_data)}")
        except Exception as e:
            print(f"[Critical Render Crash] Page {page_idx}: {e}")
        self.update_bookmark_ui()

    def toggle_pause(self):
        if not self.is_paused:
            # We are pausing
            self.is_paused = True
            self.stop_requested = True
            self.status_label.configure(text="Pausing...")
            self.pause_btn.configure(text="Resume", fg_color="#2E7D32", hover_color="#1B5E20")
        else:
            # We are resuming
            self.is_paused = False
            self.stop_requested = False
            self.pause_btn.configure(text="Pause", fg_color="#C62828", hover_color="#B71C1C")
            self.status_label.configure(text="Resuming translation...")
            self.menu_bar.entryconfigure("Export", state="disabled") 

            
            # Re-bind the click so it can be paused again
            self.pause_btn.configure(command=self.toggle_pause)
            
            start_page = self.total_pages # Start from the next un-translated page
            if self.current_pdf_path:
                threading.Thread(target=self.process_pdf, args=(self.current_pdf_path, start_page), daemon=True).start()

    def get_valid_pdf_path(self):
        if self.current_pdf_path and os.path.exists(self.current_pdf_path):
            return True
        # If missing, ask the user
        file_path = filedialog.askopenfilename(title="Select Original PDF to Re-Scan", filetypes=[("PDF Files", "*.pdf")])
        if file_path:
            self.current_pdf_path = file_path
            # optionally save to cache
            if self.all_page_data:
                self.all_page_data[0]['original_pdf_path'] = file_path
                self.save_page_to_cache(self.all_page_data[0])
            return True
        return False

    def rescan_current_page(self):
        if not self.all_page_data: return
        if not self.get_valid_pdf_path(): return
        
        self.rescan_btn.configure(state="disabled")
        if hasattr(self, 'rescan_err_btn'): self.rescan_err_btn.configure(state="disabled")
        threading.Thread(target=self.process_pdf, args=(self.current_pdf_path, self.current_page_idx, True), daemon=True).start()

    def rescan_all_errors(self):
        if not self.all_page_data: return
        if not self.get_valid_pdf_path(): return
        
        pages_to_rescan = []
        for i, pd in enumerate(self.all_page_data):
            en = pd.get('english', '')
            lit = pd.get('literal', '')
            if "[Translation Error]" in en or "[OCR Trans Error:" in en or "[Layout Error:" in en or "[Sync Error]" in lit:
                pages_to_rescan.append(i)
                
        if not pages_to_rescan:
            messagebox.showinfo("Re-Scan", "No translation errors found!")
            return
            
        self.rescan_btn.configure(state="disabled")
        self.rescan_err_btn.configure(state="disabled")
        threading.Thread(target=self.process_pdf, args=(self.current_pdf_path, 0, False, pages_to_rescan), daemon=True).start()

    def process_pdf(self, file_path, start_page=0, single_page_mode=False, pages_to_process=None):
        try:
            # CRITICAL: Align self.cache_dir with the book-specific subfolder
            # This MUST match where ocr_worker saves images (line ~184)
            book_name = os.path.basename(file_path).replace(".pdf", "")
            self.cache_dir = get_app_path(os.path.join(".translator_cache", book_name))
            os.makedirs(self.cache_dir, exist_ok=True)
            
            session_at_start = self.current_session
            self.stop_requested = False
            self.is_paused = False
            if not single_page_mode and not pages_to_process:
                self.schedule_ui_update(lambda: self.pause_btn.configure(state="normal", text="Pause", fg_color="#C62828", hover_color="#B71C1C", command=self.toggle_pause))
                self.schedule_ui_update(lambda: self.menu_bar.entryconfigure("Export", state="disabled"))
            else:
                self.schedule_ui_update(lambda: self.status_label.configure(text=f"Rescanning..."))
                
            doc = fitz.open(file_path)
            num_pages_to_process = len(doc)
            
            # --- PRE-ALLOCATE Order-Aware Slots ---
            if not pages_to_process:
                # SPRINT 33 FIX: Do not wipe existing cache if we are resuming
                if not self.all_page_data:
                    self.all_page_data = [None] * num_pages_to_process
                elif len(self.all_page_data) < num_pages_to_process:
                    self.all_page_data.extend([None] * (num_pages_to_process - len(self.all_page_data)))
                
                self.total_pages = num_pages_to_process
                self.schedule_ui_update(lambda: self.page_label.configure(text=f"Page {self.current_page_idx + 1} of {num_pages_to_process}"))
                self.schedule_ui_update(lambda: self.render_page(self.current_page_idx)) # Show current slot immediately

            selected_lang = self.lang_menu.get()
            is_auto = (selected_lang == "Auto-Detect")
            
            page_rect = doc[0].rect
            page_width = page_rect.width

            if self.active_executor:
                # Clear previous tasks if any
                pass
            
            # Ensure translation executor exists
            if not self.translation_executor:
                self.translation_executor = concurrent.futures.ThreadPoolExecutor(max_workers=5)
            
            ocr_lang_code = self.languages[selected_lang]["ocr"]
            src_lang_code = self.languages[selected_lang]["trans"]
            is_rtl = False
            
            translator_src = 'auto' if is_auto else src_lang_code
            translator = GoogleTranslator(source=translator_src, target='en')

            loop_iterable = pages_to_process if pages_to_process else range(start_page, num_pages_to_process if not single_page_mode else min(start_page + 1, num_pages_to_process))
            num_total = len(loop_iterable)

            if self.speed_var.get() and len(loop_iterable) > 1 and not single_page_mode:
                # MULTI-PROCESS PIPELINE
                self.schedule_ui_update(lambda: self.status_label.configure(text=f"Pipelined OCR (Parallel)..."))
                
                # Worker count: Balanced (50% cores) OR Eco (2 cores)
                cpu_count = os.cpu_count() or 4
                max_workers = 2 if self.eco_var.get() else max(1, cpu_count // 2)
                with concurrent.futures.ProcessPoolExecutor(max_workers=max_workers) as executor:
                    self.active_executor = cast(Any, executor)
                    _chunk_size = int(10) if len(loop_iterable) > 50 else int(5)
                    chunk_iterable: List[int] = list(loop_iterable)
                    for i in range(0, len(chunk_iterable), _chunk_size):
                        if self.stop_requested or self.current_session != session_at_start: break
                        
                        chunk = []
                        _end_idx = min(i + _chunk_size, len(chunk_iterable))
                        for idx in range(i, _end_idx):
                            chunk.append(chunk_iterable[idx])
                        
                        batch_num = (i // _chunk_size) + 1
                        total_batches = (len(chunk_iterable) + _chunk_size - 1) // _chunk_size
                        
                        futures = {executor.submit(ocr_worker, page_num, 
                                                  self.current_pdf_path, 
                                                  ocr_lang_code, is_auto, self.languages, translator_src, 
                                                  page_width, is_rtl, self.tesseract_path): page_num 
                                   for page_num in chunk}
                        
                        batch_trans_futures = []
                        
                        for f in concurrent.futures.as_completed(futures):
                            if self.stop_requested or self.current_session != session_at_start: break
                            p_num_expected = futures[f]
                            try:
                                res = f.result()
                                if res and isinstance(res, dict) and "error" not in res:
                                    p_num = int(cast(Any, res).get('page', 0))
                                    if p_num != p_num_expected:
                                        res['page'] = p_num_expected
                                        p_num = p_num_expected

                                    if p_num < len(self.all_page_data):
                                        self.all_page_data[p_num] = res
                                    self.save_page_to_cache(res)
                                    
                                    if p_num == self.current_page_idx and self.current_session == session_at_start:
                                        self.schedule_ui_update(lambda idx=p_num: self.render_page(int(idx)))
                                    
                                    # Phase 2: Kick off Translation
                                    source_text = str(res.get('original', ''))
                                    src_lang = str(res.get('translator_src', 'auto'))
                                    
                                    if len(source_text) > 5 and "[No Text" not in source_text:
                                        f_trans = self.translation_executor.submit(translate_worker, source_text, src_lang, self.selected_engine, self.deepl_key, self.openai_key, self.glossary)
                                        
                                        def on_trans_done(fut, page_idx=p_num, sess=session_at_start):
                                            try:
                                                t_res = fut.result()
                                                if "error" not in t_res and self.current_session == sess:
                                                    self.all_page_data[page_idx]['english'] = t_res['translated']
                                                    self.all_page_data[page_idx]['literal'] = t_res['literal']
                                                    self.save_page_to_cache(self.all_page_data[page_idx])
                                                    if page_idx == self.current_page_idx:
                                                        self.schedule_ui_update(lambda i=page_idx: self.render_page(int(i)))
                                                elif "error" in t_res:
                                                    self.all_page_data[page_idx]['english'] = "[Translation Failed]"
                                            except Exception as te:
                                                if page_idx < len(self.all_page_data):
                                                    self.all_page_data[page_idx]['english'] = f"[Error: {str(te)[:30]}]"
                                            finally:
                                                if self.current_session == sess:
                                                    self.schedule_ui_update(lambda i=page_idx+1: self.status_label.configure(text=f"Finished Page {i}"))
                                            
                                        f_trans.add_done_callback(on_trans_done)
                                        batch_trans_futures.append(f_trans)
                                    else:
                                        self.all_page_data[p_num]['english'] = ""
                                        self.schedule_ui_update(lambda i=p_num+1: self.status_label.configure(text=f"Finished Page {i} (Image Only)"))

                                    self.schedule_ui_update(self.update_idletasks)
                                elif res and "error" in res: raise Exception(res["error"])
                            except Exception as e:
                                print(f"[Parallel Task Failed] Page {p_num_expected}: {e}")
                                err_data = {
                                    "page": p_num_expected,
                                    "original": f"[Error on Page {p_num_expected + 1}]",
                                    "english": f"Processing failed: {str(e)}",
                                    "literal": "",
                                    "is_image": False,
                                    "is_rtl_page": False,
                                    "is_centered": False,
                                    "is_cover": False
                                }
                                if p_num_expected < len(self.all_page_data):
                                    self.all_page_data[p_num_expected] = err_data
                                    if p_num_expected == self.current_page_idx:
                                        self.schedule_ui_update(lambda v=p_num_expected: self.render_page(int(v)))

                        # OCR Batch Handover: Move ahead to keep the UI snappy
                        # We no longer "Hard-Wait" here to ensure the engine doesn't hang.
                        
            else:
                # SEQUENTIAL LOOP
                for page_num in loop_iterable:
                    if self.stop_requested or self.current_session != session_at_start: break
                    if self.is_paused:
                        while self.is_paused and not self.stop_requested: time.sleep(0.5)
                    self.process_single_page(doc, page_num, translator, ocr_lang_code, is_auto, selected_lang, is_rtl, single_page_mode, pages_to_process, num_pages_to_process)
                    # SPRINT 27: Stealth heartbeat (3.0s delay)
                    time.sleep(3.0 + random.uniform(0.1, 0.5))

            if not self.stop_requested:
                if pages_to_process:
                    self.schedule_ui_update(lambda: self.status_label.configure(text="Batch Re-scan Complete!"))
                elif single_page_mode:
                    self.schedule_ui_update(lambda: self.status_label.configure(text="Page Re-scan Complete!"))
                else:
                    self.schedule_ui_update(lambda: self.status_label.configure(text="Complete!"))
                    self.schedule_ui_update(lambda: self.pause_btn.configure(state="disabled"))
                    
                self.schedule_ui_update(lambda: self.menu_bar.entryconfigure("Export", state="normal"))
                self.schedule_ui_update(lambda: self.publish_btn.configure(state="normal"))
            
            if 'doc' in locals() and doc:
                doc.close() # Ensure closure
            
            self.active_executor = None
            import gc
            gc.collect() # Force cleanup
            
        except Exception as e:
            self.schedule_ui_update(lambda e_msg=str(e): messagebox.showerror("Error", e_msg))
        finally:
             self.schedule_ui_update(lambda: self.status_label.configure(text="Ready"))
             if single_page_mode or pages_to_process:
                 self.schedule_ui_update(lambda: self.rescan_btn.configure(state="normal"))
                 if hasattr(self, 'rescan_err_btn'): self.schedule_ui_update(lambda: self.rescan_err_btn.configure(state="normal"))

    def on_closing(self):
        """Gracefully shuts down all background tasks before exiting."""
        self.stop_requested = True
        try:
            if hasattr(self, 'ocr_executor') and self.ocr_executor:
                self.ocr_executor.shutdown(wait=False)
            if hasattr(self, 'translation_executor') and self.translation_executor:
                self.translation_executor.shutdown(wait=False)
            if hasattr(self, 'active_executor') and self.active_executor:
                self.active_executor.shutdown(wait=False)
            
            if self.supabase and hasattr(self.supabase.auth, 'close'):
                self.supabase.auth.close()
        except: pass
        self.destroy()
        sys.exit(0)

    def literal_word_by_word(self, trans, text):
        if not text: return ""
        words = text.split()
        lits = []
        for w in words[:200]: # Cap for speed
            try: lits.append(trans.translate(w))
            except: lits.append(w)
        return " ".join(lits)

    def process_single_page(self, doc, page_num, translator, ocr_lang_code, is_auto, selected_lang, is_rtl, single_page_mode, pages_to_process, num_pages_to_process):
        """Processes a single page sequentially using the ocr_worker for consistency."""
        session_at_start = self.current_session
        try:
            # Re-confirm page dimensions
            page_rect = doc[0].rect
            page_width = page_rect.width
            translator_src = translator.source
            
            # Update status
            progress_text = f"Processing Page {page_num + 1} of {num_pages_to_process}..."
            self.schedule_ui_update(lambda t=progress_text: self.status_label.configure(text=t))
            if num_pages_to_process > 0:
                self.schedule_ui_update(lambda v=(page_num + 1) / num_pages_to_process: self.progress_bar.set(v))
            
            # Phase 1: OCR Extraction (Direct call to worker for sequential safety)
            res = ocr_worker(page_num, self.current_pdf_path, ocr_lang_code, is_auto, 
                           self.languages, translator_src, page_width, is_rtl, self.tesseract_path)
            
            if res and isinstance(res, dict) and "error" not in res:
                p_num = int(res.get('page', page_num))
                if p_num < len(self.all_page_data):
                    self.all_page_data[p_num] = res
                self.save_page_to_cache(res)
                
                # Render immediately
                if self.current_session == session_at_start:
                    self.schedule_ui_update(lambda idx=p_num: self.render_page(int(idx)))
                
                # Phase 2: Translation
                source_text = str(res.get('original', ''))
                src_lang = str(res.get('translator_src', 'auto'))
                
                if len(source_text) > 5 and "[No Text" not in source_text:
                    t_res = translate_worker(source_text, src_lang, self.selected_engine, 
                                           self.deepl_key, self.openai_key, self.glossary)
                    if "error" not in t_res and self.current_session == session_at_start:
                        self.all_page_data[p_num]['english'] = t_res['translated']
                        self.all_page_data[p_num]['literal'] = t_res['literal']
                        self.save_page_to_cache(self.all_page_data[p_num])
                        self.schedule_ui_update(lambda i=p_num: self.render_page(int(i)))
                
                self.schedule_ui_update(lambda i=p_num + 1: self.status_label.configure(text=f"Finished Page {i}"))
            elif res and "error" in res:
                print(f"[Sequential] OCR Error Page {page_num}: {res['error']}")
        except Exception as e:
            print(f"[Sequential] Page {page_num} failed: {e}")

    def preprocess_for_ocr(self, img):
        # Neural net OCR works better on high contrast, sharp images
        gray = ImageOps.grayscale(img)
        # Increase contrast and sharpen to help character recognition
        enhanced = ImageOps.autocontrast(gray, cutoff=2)
        sharpened = enhanced.filter(ImageFilter.SHARPEN)
        return sharpened

    def deep_preprocess_for_ocr(self, pil_img):
        # 1. Convert PIL to OpenCV format (Grayscale)
        open_cv_image = np.array(pil_img.convert('L'))
        
        # 1.5 Upscale the image to boost clarity for tiny text before further processing
        # This acts as a magnifying glass, making tiny/fuzzy letters much clearer for Tesseract
        open_cv_image = cv2.resize(open_cv_image, None, fx=1.5, fy=1.5, interpolation=cv2.INTER_CUBIC)
        
        # 2. Border Detection & Removal (Rescue for decorative frames like Page 3)
        # We look for large contours that might be a decorative border
        _, thresh_border = cv2.threshold(open_cv_image, 128, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        contours, _ = cv2.findContours(thresh_border, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        h_img, w_img = open_cv_image.shape
        img_area = h_img * w_img
        
        for cnt in contours:
            area = cv2.contourArea(cnt)
            # If a contour covers > 60% of image with a large aspect ratio, it's likely a border
            if area > img_area * 0.6:
                x, y, w_border, h_border = cv2.boundingRect(cnt)
                # If it's near the edges, we white out the border area
                if x < 50 or y < 50 or (x + w_border) > w_img - 50:
                    # Create a mask to keep only the inside
                    mask = np.zeros_like(open_cv_image)
                    # Shrink the border rectangle slightly to keep text
                    inside_rect = (x + 30, y + 30, w_border - 60, h_border - 60)
                    cv2.rectangle(mask, (inside_rect[0], inside_rect[1]), 
                                  (inside_rect[0] + inside_rect[2], inside_rect[1] + inside_rect[3]), 255, -1)
                    # Apply mask: anything outside become white
                    open_cv_image = cv2.bitwise_and(open_cv_image, mask)
                    open_cv_image[mask == 0] = 255
                    break

        # 3. Bilateral Filtering (Preserves dots/diacritics while removing noise)
        denoised = cv2.bilateralFilter(open_cv_image, 9, 75, 75)
        
        # 3b. Stronger Contrast & Sharpening (Academic Tuning)
        # Convert back to PIL for high-quality sharpening
        temp_pil = Image.fromarray(denoised)
        # Apply high-pass sharpening to define thin characters (brackets, dots)
        temp_pil = temp_pil.filter(ImageFilter.SHARPEN)
        temp_pil = ImageOps.autocontrast(temp_pil, cutoff=1)
        # SPRINT 31: Auto-Contrast Equalization to help OCR find faint numbers
        temp_pil = ImageOps.equalize(temp_pil)
        # Back to CV
        denoised = np.array(temp_pil)

        # 4. Deskewing
        _, thresh = cv2.threshold(denoised, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)
        coords = np.column_stack(np.where(thresh > 0))
        
        if len(coords) > 0:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
        else:
            angle = 0
            
        (h, w) = denoised.shape[:2]
        center = (w // 2, h // 2)
        M = cv2.getRotationMatrix2D(center, angle, 1.0)
        rotated = cv2.warpAffine(denoised, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
        
        # 5. Hybrid Thresholding (Otsu for styled fonts if Deep Cleanup is on)
        if self.deep_cleanup_var.get():
            # Otsu is better for thick stylized fonts on uniform backgrounds (Covers)
            _, final_thresh = cv2.threshold(rotated, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        else:
            # Adaptive is better for standard pages with uneven lighting
            final_thresh = cv2.adaptiveThreshold(rotated, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                                cv2.THRESH_BINARY, 11, 2)
        
        # 6. "Quiet Zone" Padding
        final_thresh = cv2.copyMakeBorder(final_thresh, 20, 20, 20, 20, cv2.BORDER_CONSTANT, value=[255])
        
        # SPRINT 31: Speckle Filter XL (Increased threshold to 15px area)
        # Removes larger scanner noise/dust that causes "Ghosting"
        labels_info = cv2.connectedComponentsWithStats(cv2.bitwise_not(final_thresh), connectivity=8)
        num_labels, labels, stats, _ = labels_info
        for i in range(1, num_labels):
            if stats[i, cv2.CC_STAT_AREA] < 15:
                final_thresh[labels == i] = 255
                
        # SPRINT 30: Ink Density Calculation
        ink_pixels = np.sum(final_thresh == 0)
        ink_density = ink_pixels / final_thresh.size
        
        return Image.fromarray(final_thresh), ink_density

    def clean_ocr_text(self, text):
        if not text: return ""
        
        # SPRINT 33: "Clean" Toggle Logic
        # If the user turns ON the Clean switch, we aggressively hunt down page-border noise
        if hasattr(self, 'deep_cleanup_var') and self.deep_cleanup_var.get():
            # 1. Strip common isolated noise symbols and border fragments that OCR hallucinates
            # TARGET NOISE: ~ ^ | _ \ / ` 
            # We use word boundaries / spacing to ensure we don't delete math equations or inline brackets
            
            # Remove isolated symbols (surrounded by spaces)
            text = re.sub(r'(?<= )[\~\|\_\^\/\\](?= )', '', text)
            
            # 2. Remove repetitive bridge characters (borders)
            # e.g., ".......", "-------", "_____"
            text = re.sub(r'[\.\-\_\=]{4,}', ' ', text)
            
            # 3. Heuristic Filter: Remove "words" that are mostly symbols or numbers
            lines = text.split('\n')
            cleaned_lines = []
            for line in lines:
                words = line.split()
                valid_words = []
                for word in words:
                    # If word has at least 1 letter or is a common number/punc, keep it
                    letter_count = len(re.findall(r'[a-zA-Z\u00C0-\u017F\u0400-\u04FF\u0600-\u06FF\u4E00-\u9FFF]', word))
                    if letter_count > 0 or word.isnumeric() or len(word) < 4:
                        valid_words.append(word)
                
                cleaned_line = " ".join(valid_words).strip()
                # Filter lines that are just single random characters (often border noise)
                if len(cleaned_line) > 1:
                    cleaned_lines.append(cleaned_line)
                    
            text = "\n".join(cleaned_lines)
            
        return text

    def restore_arabic_numerals(self, text):
        # Maps Western (English) digits back to Eastern Arabic digits for UI aesthetics
        mapping = {'0':'٠','1':'١','2':'٢','3':'٣','4':'٤','5':'٥','6':'٦','7':'٧','8':'٨','9':'٩'}
        for w, e in mapping.items():
            text = text.replace(w, e)
        return text

    def normalize_to_western_digits(self, text):
        # Maps Eastern Arabic digits BACK to Western for the translation engine
        mapping = {'٠':'0','١':'1','٢':'2','٣':'3','٤':'4','٥':'5','٦':'6','٧':'7','٨':'8','٩':'9'}
        for e, w in mapping.items():
            text = text.replace(e, w)
        return text

    def detect_script(self, text):
        """Returns the likely Tesseract OCR code(s) joined by '+' and RTL status."""
        if not text or len(text) < 5:
            return {"ocr": "eng", "trans": "auto", "rtl": False}
            
        # Character ranges with broader detection
        arabic_chars = len(re.findall(r'[\u0600-\u06FF]', text))
        hebrew_chars = len(re.findall(r'[\u0590-\u05FF]', text))
        greek_chars = len(re.findall(r'[\u0370-\u03FF]', text))
        cyrillic_chars = len(re.findall(r'[\u0400-\u04FF]', text))
        chinese_chars = len(re.findall(r'[\u4E00-\u9FFF]', text))
        latin_chars = len(re.findall(r'[a-zA-Z\u00C0-\u017F]', text))
        
        counts = {
            "ara": arabic_chars,
            "heb": hebrew_chars,
            "ell": greek_chars,
            "rus": cyrillic_chars,
            "chi_sim": chinese_chars,
            "eng": latin_chars
        }
        
        # Mapping back to translator codes for single-language fallbacks
        trans_map = {
            "ara": "ar", "heb": "he", "ell": "el", "rus": "ru", "chi_sim": "zh-CN", "eng": "en", "fra": "fr"
        }
        
        total_chars = sum(counts.values())
        if total_chars == 0:
            return {"ocr": "eng", "trans": "auto", "rtl": False}
            
        # Identify all significant scripts (> 2% of characters or > 2 chars total)
        detected_ocr_codes = []
        is_rtl = False
        
        # Sort to keep eng/latin/fra at the end (Tesseract convention)
        potential_codes = ["ara", "heb", "ell", "rus", "chi_sim", "fra", "eng"]
        
        # Priority Hack: If any latin is found, use 'fra' instead of 'eng' as it handles accents better
        if counts["eng"] > 0:
            counts["fra"] = counts["eng"]
            counts["eng"] = 0

        for code in potential_codes:
            char_count = counts.get(code, 0)
            if char_count > 2 or (char_count > 0 and char_count / total_chars > 0.02):
                if code == "fra":
                    # Always include eng with fra for safety
                    detected_ocr_codes.extend(["fra", "eng"])
                else:
                    detected_ocr_codes.append(code)
                
                if code in ["ara", "heb"]:
                    is_rtl = True
        
        if not detected_ocr_codes:
            detected_ocr_codes = ["eng"]
            
        final_ocr = "+".join(detected_ocr_codes)
        
        # Determine likely primary source language for the translator
        # If multiple, use 'auto'
        if len(detected_ocr_codes) > 1:
            final_trans = "auto"
        else:
            final_trans = trans_map.get(detected_ocr_codes[0], "auto")
            
        return {
            "ocr": final_ocr,
            "trans": final_trans,
            "rtl": is_rtl
        }

    def robust_translate(self, translator, text):
        """Attempts to translate text with exponential backoff, length safety, and stealth diagnostics."""
        if not text: return ""
        
        # SPRINT 29: Normalize numerals to Western for Google accuracy
        text = self.normalize_to_western_digits(text)
        
        # SPRINT 31: Unicode Isolation Shield (LRI/PDI markers)
        text = re.sub(r'(\d+[\-/]\d+[\-/]?[0-9]*[\-/]?[0-9]*)', u' \u2066\\1\u2069 ', text) 
        
        text = self.rebalance_scripts(text)
        
        if len(text) < 3: return text
        
        # SPRINT 27: Tight Character Limit (2,500 chars) for maximum stealth
        if len(text) > 2500:
            parts = [text[i:i+2500] for i in range(0, len(text), 2500)]
            results = [self.robust_translate(translator, p) for p in parts]
            return "\n".join(results)

        import time
        import random
        # 3-tier backoff for total reliability
        for attempt in range(3):
            try:
                translated = "[Translation Error]"
                if self.selected_engine == "DeepL" and self.deepl_key:
                    import deepl
                    t = deepl.Translator(self.deepl_key)
                    src = None if translator.source == 'auto' else translator.source.upper()
                    res = t.translate_text(text, target_lang="EN-US", source_lang=src)
                    translated = res.text
                elif self.selected_engine == "GPT-4o" and self.openai_key:
                    from openai import OpenAI
                    client = OpenAI(api_key=self.openai_key)
                    
                    gloss_instr = ""
                    if self.glossary:
                        gloss_list = [f"'{k}' -> '{v}'" for k, v in self.glossary.items()]
                        gloss_instr = "IMPORTANT: Use the following custom terminology glossary: " + ", ".join(gloss_list) + "."
                        
                    response = client.chat.completions.create(
                      model="gpt-4o",
                      messages=[
                        {"role": "system", "content": f"You are a professional translator. Translate the following text from {translator.source} to English. Preserve formatting and nuances. Output ONLY the translated text. {gloss_instr}"},
                        {"role": "user", "content": text}
                      ]
                    )
                    translated = response.choices[0].message.content.strip()
                else: # Default Google
                    # Pre-processing Glossary for Google
                    if self.glossary:
                        for orig, target in self.glossary.items():
                            pattern = re.compile(re.escape(orig), re.IGNORECASE)
                            text = pattern.sub(target, text)
                            
                    fresh_translator = GoogleTranslator(source=translator.source, target=translator.target)
                    translated = str(fresh_translator.translate(text))
                    
                    # Passthrough check
                    if translated.strip().lower() == text.strip().lower() and len(text) > 15:
                        auto_translator = GoogleTranslator(source='auto', target='en')
                        translated = str(auto_translator.translate(text))
                
                return translated
            except Exception as e:
                print(f"DEBUG: Translation Attempt {attempt+1} failed ({len(text)} chars): {str(e)}")
                time.sleep(2 ** attempt + random.uniform(0.1, 0.5))
                
        return "[Translation Failed]"

    def rebalance_scripts(self, text):
        """Fixes common French/Greek script confusion and bracket mangling."""
        if not text: return ""
        
        # 1. Fix French 'le' becoming Greek 'ἰο' or 'ἰο'
        # Tesseract often confuses 'le' with 'ἰo' or 'ἰο' (Greek iota-omicron)
        # We look for Greek snippets that are likely French words
        text = re.sub(r'\bἰ[οo]\b', 'le', text)
        text = re.sub(r'\bἰ[δθ]\b', 'le', text) # Alternative mangling
        
        # 2. Fix Greek 'αο' becoming French 'que' or 'ce'
        # Scholarship text: 'αο' (Greek alpha-omicron) is often 'que' or 'ce' or 'ou'
        # If surrounded by French latin text, these are almost certainly errors
        # (This is a subtle fix, we only apply to very common small errors)
        text = re.sub(r' ([αα]ο|οα) ', ' le ', text)
        
        # 3. Bracket mangling cleanup
        # Fix (( into (
        text = re.sub(r'\({2,}', '(', text)
        text = re.sub(r'\){2,}', ')', text)
        # Fix ((. into (f. (Common for folio markers in intro)
        text = re.sub(r'\(\(\. ', '(f. ', text)
        # Fix 6"-7 to 6-7
        text = re.sub(r'6\"-7', '6-7', text)
        
        # 4. Fix Greek iota mangling in French words
        text = re.sub(r'\bἰο\b', 'le', text)
        text = re.sub(r'ἰο tome', 'le tome', text)
        
        return text

    def tokenize_text(self, text: str) -> list:
        # Unified regex for word/number detection
        return re.findall(r'\b\w+\b', str(text))

    def populate_ui_skeleton(self, page_num, original, translated, literal, is_image, is_rtl_override=None, is_centered=False, is_cover=False, cover_image=None):
        """Updates the permanent UI skeleton with the current page data instantly."""
        # Pic 2/3 Clean Look: Hide "No Text" labels
        display_orig = "" if "[No Text Detected]" in str(original) else original
        display_trans = "" if "[No Text Detected]" in str(translated) or "[Translation Pending]" in str(translated) else translated
        if "[Translation Pending]" in str(translated): display_trans = "[Translation Pending]"

        for box in [self.orig_text, self.trans_text, self.lit_text]:
            box.configure(state="normal")
            box.delete("1.0", "end")
            box.configure(state="disabled")
        self.visual_img_label.configure(image="")

        is_rtl = is_rtl_override if is_rtl_override is not None else (self.lang_menu.get() in ["Arabic", "Hebrew", "Yiddish"])
        
        # UI is 1-based, internal is 0-based
        self.page_number_label.configure(text=f"Page {page_num + 1}")

        # Pic 1: Uniformity. All pages now show the 3-column layout.
        show_text_panels = True 
        
        if show_text_panels:
            self.orig_container.grid(row=1, column=1, sticky="nsew")
            self.trans_container.grid(row=1, column=3, sticky="nsew")
            self.mid_container.grid(row=1, column=2, columnspan=1, sticky="nsew")
            
            # Transcription Column
            if is_rtl:
                self.orig_text.tag_config("rtl", justify="right")
                self.insert_and_tag_words(self.orig_text, display_orig, is_rtl=True)
            else:
                self.insert_and_tag_words(self.orig_text, display_orig)
                
            # Translation Column
            self.insert_and_tag_words(self.trans_text, display_trans, is_centered=is_centered)
        else:
            # ONLY Page 0 (Cover) shows as Image Only
            self.orig_container.grid_remove()
            self.trans_container.grid_remove()
            self.mid_container.grid(row=1, column=0, columnspan=4, sticky="nsew")

        # Middle Column (1-1 Match OR Image feed)
        if self.one_one_var.get() and show_text_panels:
            self.visual_img_label.pack_forget()
            self.lit_text.pack(padx=15, pady=15, fill="both", expand=True)
            self.insert_and_tag_words(self.lit_text, literal)
        else:
            self.lit_text.pack_forget()
            self.visual_img_label.pack(padx=5, pady=25, fill="both", expand=True, anchor="n")
            
            img_to_show = cover_image
            if not img_to_show and isinstance(original, Image.Image):
                img_to_show = original
            
            if img_to_show:
                display_img = img_to_show.copy()
                max_w, max_h = (500, 800)
                display_img.thumbnail((max_w, max_h))
                tk_img = ctk.CTkImage(light_image=display_img, dark_image=display_img, size=(display_img.width, display_img.height))
                self.visual_img_label.configure(image=tk_img, text="")
                self.visual_img_label._image = tk_img # Keep reference
            else:
                self.visual_img_label.configure(image=None, text="[No Image Capture]")

    def insert_and_tag_words(self, textbox, text, is_rtl=False, is_centered=False):
        textbox.configure(state="normal")
        textbox.delete("1.0", "end")
        
        # Simple insertion without word-level highlighting for clarity
        clean_text = str(text)
        textbox.insert("1.0", clean_text)
        
        if is_rtl:
            textbox.tag_add("content", "1.0", "end")
            textbox.tag_config("content", justify="right")
        elif is_centered:
            textbox.tag_add("content", "1.0", "end")
            textbox.tag_config("content", justify="center")

    def link_highlighting(self, textboxes):
        pass # Highlighting disabled as requested

    def save_as_pdf(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
        if not file_path: return
        
        try:
            doc = SimpleDocTemplate(file_path, pagesize=LETTER, 
                                    rightMargin=50, leftMargin=50, 
                                    topMargin=50, bottomMargin=50)
            
            styles = getSampleStyleSheet()
            
            # Create custom Styles
            title_style = ParagraphStyle(
                'TitleStyle', parent=styles['Heading1'],
                fontName='Helvetica-Bold', fontSize=24, spaceAfter=20, alignment=1
            )
            
            column_style = ParagraphStyle(
                'ColumnStyle', parent=styles['Normal'],
                fontName='Helvetica', fontSize=10, leading=12
            )
            
            # If the user has a custom font loaded, try to use it
            try:
                # We assume self.font_family is a standard PDF font or one we've registered
                # For more robust Unicode, we'd need to properly register TTFs from disk
                column_style.fontName = self.font_family
            except:
                pass
                
            story = []
            
            # Add Title
            title_text = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Translation"
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 0.5 * inch))
            
            for data in self.all_page_data:
                # Page Header
                story.append(Paragraph(f"Page {data['page'] + 1}", styles['Heading2']))
                story.append(Spacer(1, 0.1 * inch))
                
                # If page is an image, just add the image
                if data.get('is_image') and data.get('cover_image_path') and os.path.exists(data['cover_image_path']):
                    img = RLImage(data['cover_image_path'], width=6 * inch, height=4 * inch, kind='proportional')
                    story.append(img)
                    story.append(Spacer(1, 0.3 * inch))
                else:
                    # Side-by-side Table
                    orig_text = str(data.get('original', '')).replace('\n', '<br/>')
                    trans_text = str(data.get('english', '')).replace('\n', '<br/>')
                    
                    p_orig = Paragraph(orig_text, column_style)
                    p_trans = Paragraph(trans_text, column_style)
                    
                    table_data = [[p_orig, p_trans]]
                    # Use 45% width for each column to leave room for the middle gutter
                    t = Table(table_data, colWidths=[3 * inch, 3 * inch])
                    
                    t.setStyle(TableStyle([
                        ('VALIGN', (0,0), (-1,-1), 'TOP'),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
                        ('BOTTOMPADDING', (0,0), (-1,-1), 12),
                        ('TOPPADDING', (0,0), (-1,-1), 12),
                        ('BACKGROUND', (1,0), (1,0), colors.whitesmoke) # Subtle highlight for translation
                    ]))
                    
                    story.append(t)
                    story.append(Spacer(1, 0.3 * inch))

            doc.build(story)
            messagebox.showinfo("Success", "Professional Bilingual PDF exported successfully!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to build PDF: {str(e)}")

    def handle_export(self, choice):
        if choice == "Export PDF":
            self.save_as_pdf()
        elif choice == "Export Word":
            self.save_as_docx()
        elif choice == "Export ePub":
            self.save_as_epub()

    def save_as_docx(self):
        """Exports the Bilingual translation to a Word document."""
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            messagebox.showwarning("Missing Library", "Please install 'python-docx' to use this feature.\nRun: pip install python-docx")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Files", "*.docx")])
        if not file_path: return

        try:
            doc = Document()
            title = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Translation"
            doc.add_heading(title, 0)

            for data in self.all_page_data:
                doc.add_heading(f"Page {data['page'] + 1}", level=1)
                
                # Table for side-by-side
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                cells = table.rows[0].cells
                
                cells[0].text = data.get('original', '')
                cells[1].text = data.get('english', '')
                
                # Image if it exists
                if data.get('is_image') and data.get('cover_image_path') and os.path.exists(data['cover_image_path']):
                    doc.add_picture(data['cover_image_path'], width=Inches(4))

            doc.save(file_path)
            messagebox.showinfo("Success", "Word document exported successfully!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to save Word doc: {str(e)}")

    def save_as_epub(self):
        """Exports the translated book to ePub format."""
        try:
            from ebooklib import epub
        except ImportError:
            messagebox.showwarning("Missing Library", "Please install 'ebooklib' and 'lxml' to use this feature.\nRun: pip install ebooklib lxml")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".epub", filetypes=[("ePub Files", "*.epub")])
        if not file_path: return

        try:
            book = epub.EpubBook()
            title = os.path.basename(self.current_pdf_path) if self.current_pdf_path else "Bilingual Book"
            book.set_title(title)
            book.set_language('en')
            
            chapters = []
            for data in self.all_page_data:
                p_num = data['page'] + 1
                c = epub.EpubHtml(title=f"Page {p_num}", file_name=f"page_{p_num}.xhtml", lang='en')
                content = f"<h1>Page {p_num}</h1>"
                
                # Add image if exists
                if data.get('is_image') and data.get('cover_image_path') and os.path.exists(data['cover_image_path']):
                    # We would need to add the item to the book properly for real images
                    content += f"<p>[Image Page]</p>"
                
                orig = data.get('original', '').replace('\n', '<br/>')
                trans = data.get('english', '').replace('\n', '<br/>')
                content += f"<h3>Original</h3><p>{orig}</p><h3>Translation</h3><p>{trans}</p>"
                
                c.content = content
                book.add_item(c)
                chapters.append(c)

            book.toc = tuple(chapters)
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            book.spine = ['nav'] + chapters
            
            epub.write_epub(file_path, book, {})
            messagebox.showinfo("Success", "ePub exported successfully!")
        except Exception as e:
            messagebox.showerror("Export Error", f"Failed to save ePub: {str(e)}")

    def schedule_ui_update(self, func):
        self.after(0, func)

if __name__ == "__main__":
    app = PDFTranslatorApp()
    app.mainloop()
